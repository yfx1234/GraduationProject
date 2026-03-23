from __future__ import annotations

import json
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(r'D:\Xstarlab\UEProjects\GraduationProject\GraduationProject')
OUT_DIR = ROOT / 'output' / 'doc'
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / '文献综述-终稿版-低空小目标视觉检测跟踪与拦截控制研究综述.docx'
META_PATH = ROOT / 'tmp' / 'docs' / 'ref_meta.json'

TITLE = '低空小目标视觉检测、跟踪与拦截控制研究综述'
AUTHOR_LINE = '作者：颜飞翔    学号：2229409056'
DATE_LINE = '完成日期：2026年3月21日'

meta = json.loads(META_PATH.read_text(encoding='utf-8'))

fallback = {
    1: ('Feature Pyramid Networks for Object Detection', 'CVPR', '2017', '10.1109/CVPR.2017.106'),
    2: ('Focal Loss for Dense Object Detection', 'ICCV', '2017', '10.1109/ICCV.2017.324'),
    3: ('EfficientDet: Scalable and Efficient Object Detection', 'CVPR', '2020', '10.1109/CVPR42600.2020.01079'),
    4: ('End-to-End Object Detection with Transformers', 'ECCV', '2020', '10.1007/978-3-030-58452-8_13'),
    5: ('QueryDet: Cascaded Sparse Query for Accelerating High-Resolution Small Object Detection', 'CVPR', '2022', '10.1109/CVPR52688.2022.01330'),
    6: ('Small Object Detection in Unmanned Aerial Vehicle Images Using Feature Fusion and Scaling-Based Single Shot Detector With Spatial Context Analysis', 'IEEE TCSVT', '2019', '10.1109/TCSVT.2019.2905881'),
    7: ('Focus-and-Detect: A Small Object Detection Framework for Aerial Images', 'Signal Processing: Image Communication', '2022', '10.1016/j.image.2022.116675'),
    8: ('MFFSODNet: Multiscale Feature Fusion Small Object Detection Network for UAV Aerial Images', 'IEEE TIM', '2024', '10.1109/TIM.2024.3381272'),
    9: ('A Context-Scale-Aware Detector and a New Benchmark for Remote Sensing Small Weak Object Detection in Unmanned Aerial Vehicle Images', 'JAG', '2022', '10.1016/j.jag.2022.102966'),
    10: ('Small Object Detection in Unmanned Aerial Vehicle Images Using Multi-Scale Hybrid Attention', 'Engineering Applications of Artificial Intelligence', '2023', '10.1016/j.engappai.2023.107455'),
    11: ('Tiny Object Detection in Aerial Images', 'ICPR', '2021', '10.1109/ICPR48806.2021.9413340'),
    12: ('Detection and Tracking Meet Drones Challenge', 'IEEE TPAMI', '2021', '10.1109/TPAMI.2021.3119563'),
    13: ('The Unmanned Aerial Vehicle Benchmark: Object Detection, Tracking and Baseline', 'IJCV', '2019', '10.1007/s11263-019-01266-1'),
    14: ('Vision-Based Anti-UAV Detection and Tracking', 'IEEE TITS', '2022', '10.1109/TITS.2022.3177627'),
    15: ('Anti-UAV: A Large-Scale Benchmark for Vision-Based UAV Tracking', 'IEEE TMM', '2021', '10.1109/TMM.2021.3128047'),
    16: ('A Unified Transformer-Based Tracker for Anti-UAV Tracking', 'CVPRW', '2023', '10.1109/CVPRW59228.2023.00305'),
    17: ('SeaDronesSee: A Maritime Benchmark for Detecting Humans in Open Water', 'WACV', '2022', '10.1109/WACV51458.2022.00374'),
    18: ('BIRDSAI: A Dataset for Detection and Tracking in Aerial Thermal Infrared Videos', 'WACV', '2020', '10.1109/WACV45572.2020.9093284'),
    19: ('WebUAV-3M: A Benchmark for Unveiling the Power of Million-Scale Deep UAV Tracking', 'IEEE TPAMI', '2022', '10.1109/TPAMI.2022.3232854'),
    20: ('Unmanned Aerial Vehicle Visual Detection and Tracking Using Deep Neural Networks: A Performance Benchmark', 'ICCVW', '2021', '10.1109/ICCVW54120.2021.00142'),
    21: ('VisDrone-DET2019: The Vision Meets Drone Object Detection in Image Challenge Results', 'ICCVW', '2019', '10.1109/ICCVW.2019.00030'),
    22: ('VisDrone-DET2021: The Vision Meets Drone Object Detection Challenge Results', 'ICCVW', '2021', '10.1109/ICCVW54120.2021.00319'),
    23: ('On the Detection of Unauthorized Drones - Techniques and Future Perspectives: A Review', 'IEEE Sensors Journal', '2022', '10.1109/JSEN.2022.3171293'),
    24: ('Simple Online and Realtime Tracking with a Deep Association Metric', 'ICIP', '2017', '10.1109/ICIP.2017.8296962'),
    25: ('ByteTrack: Multi-Object Tracking by Associating Every Detection Box', 'ECCV', '2022', '10.1007/978-3-031-20047-2_1'),
    26: ('Tracking Objects as Points', 'ECCV', '2020', '10.1007/978-3-030-58548-8_28'),
    27: ('SiamRPN++: Evolution of Siamese Visual Tracking With Very Deep Networks', 'CVPR', '2019', '10.1109/CVPR.2019.00441'),
    28: ('Learning Discriminative Model Prediction for Tracking', 'ICCV', '2019', '10.1109/ICCV.2019.00628'),
    29: ('Improving Multiple Object Tracking with Single Object Tracking', 'CVPR', '2021', '10.1109/CVPR46437.2021.00248'),
    30: ('Observation-Centric SORT: Rethinking SORT for Robust Multi-Object Tracking', 'CVPR', '2023', '10.1109/CVPR52729.2023.00934'),
    31: ('HOTA: A Higher Order Metric for Evaluating Multi-Object Tracking', 'IJCV', '2020', '10.1007/s11263-020-01375-2'),
    32: ('Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs', 'IEEE/ASME TMech', '2019', '10.1109/TMECH.2019.2906430'),
    33: ('Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets', 'IEEE RA-L', '2023', '10.1109/LRA.2023.3245416'),
    34: ('Image-Based Visual Servoing of Rotorcrafts to Planar Visual Targets of Arbitrary Orientation', 'IEEE RA-L', '2021', '10.1109/LRA.2021.3101878'),
    35: ('Position-Based Visual Servoing for Target Tracking by a Quadrotor UAV', 'AIAA GNC', '2016', '10.2514/6.2016-2092'),
    36: ('Autonomous Drone Hunter Operating by Deep Learning and All-Onboard Computations in GPS-Denied Environments', 'PLOS ONE', '2019', '10.1371/journal.pone.0225092'),
    37: ('Autonomous Target Tracking of UAV Using High-Speed Visual Feedback', 'Applied Sciences', '2019', '10.3390/app9214552'),
    38: ('Robust Visual Servoing Formation Tracking Control for Quadrotor UAV Team', 'Aerospace Science and Technology', '2020', '10.1016/j.ast.2020.106061'),
    39: ('UAV Guidance: A Stereo-Based Technique for Interception of Stationary or Moving Targets', 'LNCS', '2015', '10.1007/978-3-319-22416-9_30'),
    40: ('AirSim: High-Fidelity Visual and Physical Simulation for Autonomous Vehicles', 'Springer FSR', '2017', '10.1007/978-3-319-67361-5_40'),
}


def set_font(run, east='宋体', west='Times New Roman', size=12, bold=False, superscript=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.superscript = superscript


CITATION_RE = re.compile(r'(?:\[\d+\](?:-\[\d+\])?)+')


def add_inline_text(paragraph, text, east='宋体', west='Times New Roman', size=12, citation_as_superscript=False):
    pos = 0
    for match in CITATION_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, east=east, west=west, size=size)
        run = paragraph.add_run(match.group(0))
        citation_size = 10.5 if citation_as_superscript else size
        set_font(run, east=east, west=west, size=citation_size, superscript=citation_as_superscript)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, east=east, west=west, size=size)

def config(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)


def title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, east='黑体', size=size, bold=True)


def center(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, east='宋体', size=size)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, east='黑体', size=14 if level == 1 else 12, bold=True)


def para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_text(p, text, east='宋体', size=12, citation_as_superscript=True)

def ref_para(doc, idx, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    mark = p.add_run(f'[{idx}] ')
    set_font(mark, east='宋体', size=10.5, bold=True)
    add_inline_text(p, text, east='宋体', size=10.5, citation_as_superscript=False)

def format_authors(authors):
    if not authors:
        return ''
    if len(authors) <= 3:
        return ', '.join(authors)
    return ', '.join(authors[:3]) + ', et al.'


def build_ref_text(idx):
    fb_title, fb_venue, fb_year, fb_doi = fallback[idx]
    m = meta.get(str(idx), {})
    authors = format_authors(m.get('authors', []))
    title = m.get('title') or fb_title
    venue = m.get('container') or fb_venue
    year = m.get('year') or fb_year
    doi = m.get('doi') or fb_doi
    if authors:
        return f'{authors}. {title}. {venue}, {year}. DOI: {doi}'
    return f'{title}. {venue}, {year}. DOI: {doi}'


def add_paragraphs(doc, texts):
    for t in texts:
        para(doc, t)


def build():
    doc = Document()
    config(doc)
    title(doc, TITLE)
    center(doc, AUTHOR_LINE)
    center(doc, DATE_LINE)

    heading(doc, '摘要')
    para(doc, '低空小目标兼具目标尺寸小、背景干扰强、运动状态变化快和观测条件复杂等特征，是机器视觉领域中兼具理论难度与工程价值的一类典型问题。围绕该问题，相关研究已经从早期的小目标检测模型改进逐步扩展到数据集与评测基准构建、时序跟踪与状态估计、视觉伺服与自主追踪控制以及仿真平台验证等多个层面。现有成果一方面显著提升了无人机和反无人机场景中的目标感知能力，另一方面也暴露出尺度敏感性强、时序稳定性不足、感知与控制耦合不紧以及系统级评价缺失等共性问题。总体来看，多尺度特征建模、上下文增强、低置信度观测利用、图像误差闭环控制以及仿真到真实环境迁移，构成了该方向的核心研究脉络。面向低慢小目标拦截任务，未来研究需要在检测、预测和控制的协同设计上进一步突破，并构建更加贴近任务需求的评价体系。')
    heading(doc, '关键词')
    para(doc, '低空小目标；目标检测；视觉跟踪；状态估计；视觉伺服；拦截控制；无人机视觉', indent=False)


    heading(doc, 'Abstract')
    para(doc, 'Low-altitude small targets are characterized by extremely limited pixel occupancy, strong background clutter, fast state variation, and unstable observation conditions, which makes them one of the most challenging problems in machine vision for aerial perception. Existing studies have gradually extended from small-object detection to benchmark construction, temporal tracking and state estimation, visual servoing, autonomous pursuit, interception control, and high-fidelity simulation. Recent literature shows that multiscale representation, contextual enhancement, confidence-aware association, short-term motion prediction, image-based closed-loop control, and simulation-to-real transfer have become the major technical threads in this field. At the same time, common limitations remain evident, including sensitivity to scale variation, unstable temporal continuity, insufficient coupling between perception and control, and the lack of task-oriented system evaluation. For low-altitude interception tasks, future research should place greater emphasis on the coordinated design of detection, prediction, and control, while developing evaluation criteria that better reflect closed-loop usability and engineering deployment requirements.')
    heading(doc, 'Key Words')
    para(doc, 'low-altitude small target; object detection; visual tracking; state estimation; visual servoing; interception control; UAV vision', indent=False)
    heading(doc, '1 引言')
    add_paragraphs(doc, [
        '低空目标感知问题近几年持续升温，其中最具挑战性的对象之一便是低慢小目标。所谓“低慢小”，并不只是三个工程描述词的简单叠加，而是对应着视觉系统中三个彼此耦合的困难来源：飞行高度低意味着背景中包含建筑物、树木、地平线、云层和地面纹理等大量干扰信息；速度慢意味着目标常常长时间停留在复杂背景附近，容易与背景产生混淆；尺寸小则直接决定了图像中目标所占像素极少，导致其边缘、纹理和形状信息都不稳定。正是这些因素共同作用，使得该类目标成为检测、跟踪和控制任务中的典型难例。',
        '如果仅从单帧视觉识别角度来看，低空小目标任务似乎可以被归结为一种特殊的小目标检测问题；但从实际系统需求出发，这一问题远不止于检测本身。对于动态目标而言，检测结果能否持续稳定地输出，目标一旦短时消失是否还能维持状态估计，图像误差是否可以进一步转化为可执行的控制量，控制过程中目标是否始终处在相机有效视场之内，这些问题都会直接影响系统是否真正可用。也正因如此，近年来的相关研究已经逐渐从单点算法改进走向跨模块协同。',
        '现有研究大体形成了几条相互关联的技术路线。第一条路线聚焦于目标检测本身，围绕多尺度特征融合、密集分类损失、注意力增强以及高分辨率稀疏推理展开，以提高小目标在复杂背景中的可分辨性[1]-[11]。第二条路线围绕数据集与 benchmark 建设，试图通过 UAVDT、VisDrone、Anti-UAV、WebUAV-3M 等数据集与挑战赛结果统一任务边界和评测标准[12]-[23]。第三条路线聚焦时序建模，在 Deep SORT、ByteTrack、CenterTrack、DiMP、OC-SORT 等方法中分别探索在线关联、端到端跟踪和运动状态估计问题[24]-[31]。第四条路线则进一步面向控制与系统闭环，利用视觉伺服、目标追踪和自主拦截研究把感知结果引入飞行控制回路[32]-[40]。',
        '对于以低慢小目标拦截为目标的研究而言，这几条技术路线并不是彼此孤立的。检测精度决定了后续控制是否有可靠观测，时序稳定性决定了控制器能否在漏检条件下保持有效输出，视觉伺服策略又会反过来影响目标在图像中的位置和可见性。换言之，低空小目标研究已经不再适合被简单理解为“先做检测，再做控制”的线性流程，而更应被看作一个贯穿感知、预测与决策的系统问题。',
        '围绕这一任务背景，低空小目标研究逐渐形成了从前端感知到闭环控制的连续技术链条。检测方法决定了系统能否在复杂背景中稳定发现目标，跟踪与状态估计决定了视觉观测能否在时间维度上保持连续，而视觉伺服与拦截控制则进一步决定了这些观测能否被转化为有效动作。各部分虽然关注点不同，但它们共同塑造了该方向的研究面貌。',
    ])

    heading(doc, '2 低空小目标视觉检测研究')
    heading(doc, '2.1 通用检测框架对低空小目标问题的奠基作用', level=2)
    add_paragraphs(doc, [
        '低空小目标检测虽然具有鲜明场景特征，但它的很多关键思想来自通用检测框架对小尺度目标的长期探索。FPN[1]通过构建自顶向下的特征金字塔，把高层语义信息传递到高分辨率特征图，为小目标提供了更加稳定的表征基础。对于低空场景而言，这一思想非常关键，因为尺寸很小的飞行目标往往只能在浅层特征图上保留有限细节，如果缺乏高层语义辅助，模型很容易把目标误判为背景噪声。',
        'Focal Loss[2]则从损失函数层面解决了密集检测中的类别不平衡问题。低空小目标任务中前景像素占比极低，背景区域又高度复杂，若仍然采用普通交叉熵，模型在训练时往往会过度关注大量容易分类的背景样本，导致真正难检的小目标被淹没。Focal Loss 对难样本和稀有正样本的加权强化，为后续许多小目标检测器提供了直接参考。',
        '在效率与部署层面，EfficientDet[3]展示了另一类值得注意的思路。它并没有通过无节制增加网络深度来换取指标，而是利用更合理的双向特征融合和尺度缩放策略，在精度和速度之间取得平衡。对于机载或边缘计算环境来说，这种平衡本身就具有方法论意义，因为小目标研究最终往往不是离线竞赛，而是需要进入实时系统。',
        'DETR[4]以及后续基于 Transformer 的检测思想，则把研究视角从“多尺度卷积如何覆盖目标”进一步扩展为“全局关系如何参与目标建模”。虽然 DETR 在小目标任务上存在收敛慢和细粒度表征不足的问题，但它提供了端到端统一建模的方向。QueryDet[5]则更进一步地指出，在高分辨率场景下，并不是所有位置都需要密集推理，真正的小目标常常具有稀疏分布特征，因此利用查询机制把计算集中到潜在目标区域，可能比一味提高输入分辨率更有效。',
        '综合来看，这一组通用检测研究并没有直接针对低空小目标场景给出完整解决方案，却从特征融合、损失设计、尺度建模和推理策略等多个层面奠定了方法基础。后续无人机场景研究之所以能够快速发展，很大程度上正是建立在这些基础思想之上。许多面向低空飞行目标的改进，本质上都是在这些框架内重新分配特征分辨率、上下文范围与推理成本，以适应目标像素稀缺和背景干扰强的实际条件。',
    ])

    heading(doc, '2.2 航拍与低空视角下的小目标专用检测改进', level=2)
    add_paragraphs(doc, [
        '与通用场景相比，航拍和低空视角中的小目标检测更容易受到平台运动和背景结构的影响。无人机相机自身在飞行过程中会产生姿态变化，背景纹理也会随着视角快速切换，导致目标常常缺乏稳定参照系。在这种情况下，单纯依赖局部卷积特征往往不足以分离目标与背景，模型需要更多上下文信息来判断一个小区域究竟是目标、背景纹理还是噪声。空间上下文分析与尺度增强策略[6]正是在这一背景下提出，说明小目标检测不能仅盯着目标框内部，而必须结合目标周围环境共同判断。',
        'Focus-and-Detect[7]体现了另一种解决方式，即先在全局图像中筛除大面积无关背景，再对潜在目标区域进行局部精细检测。该方法的核心价值在于承认了一个现实：在航拍大视场图像中，真正包含目标信息的区域往往非常有限，直接在整张图上均匀计算不但成本高，而且容易被复杂背景牵着走。因此，把检测问题分解为“区域选择 + 局部增强”两步，可以更有针对性地处理低空小目标。',
        'MFFSODNet[8]和多尺度混合注意力方法[10]进一步延续了“增强有用细节，抑制无关背景”的路线。前者强调浅层纹理和深层语义的多尺度协同，后者则试图通过注意力机制主动提升目标区域的响应强度。它们共同说明，航拍场景的小目标检测并非单一结构技巧可以解决，而是需要多种信息源联合发挥作用。',
        '面向弱小目标的新 benchmark[9]以及对 tiny object detection 的系统分析[11]，也从数据与任务定义层面强化了这一认识。对低空小目标而言，问题的关键不只是目标尺寸小，还包括目标与背景对比度低、存在视角畸变、长宽比不稳定，以及成像条件与距离变化带来的表观漂移。换句话说，所谓“小目标检测”在低空场景下其实是一个集合性问题，它同时包含了尺度问题、背景问题和动态问题。',
        '这类研究对后续工作最重要的启示，是必须把检测器结构、训练数据组织和场景条件放在一起考虑。脱离具体场景去讨论检测器优劣往往意义有限，因为同一个模型在一般图像上的表现，并不能直接外推到低空复杂背景中的小目标任务。尤其在飞行器目标占画面比例极低时，训练样本的尺度分布、背景多样性和标注边界一致性，往往会与网络结构同样重要。',
    ])

    heading(doc, '2.3 精度、实时性与工程可部署性的平衡', level=2)
    add_paragraphs(doc, [
        '从已有研究看，小目标检测方向的另一个长期矛盾在于精度与实时性的平衡。一方面，多尺度特征融合、上下文增强和注意力模块确实能够显著提升指标，尤其是在极小目标或复杂背景下更为明显；另一方面，这类改进往往伴随更高的参数量和更长的推理时间。当任务目标是离线测试时，这种代价尚可接受，但当检测结果需要直接进入飞行控制回路时，任何额外延迟都会被放大为系统响应迟滞。',
        '因此，当前研究中越来越多的工作开始关注“足够好的实时检测”而非“绝对最优精度”。例如 EfficientDet[3]和 QueryDet[5]都在不同意义上体现出这一趋势：前者强调网络缩放和融合效率，后者强调将计算资源集中到关键区域。这说明低空小目标检测不应简单理解为追求更高 AP，而应当以任务需求为约束，在可用性和性能之间寻找平衡。',
        '对于面向拦截的应用场景而言，这种平衡尤为重要。检测器若偶尔漏检，后端还可以通过状态估计进行短时补偿；但若检测器整体延迟过大，即便单帧精度较高，也很难支撑动态控制。因此，对该方向的评价不能仅停留在检测指标层面，而应同时考虑时间性能和闭环适应性。',
    ])


    heading(doc, '2.4 误检抑制、置信度利用与检测结果可用性', level=2)
    add_paragraphs(doc, [
        '在低空小目标场景中，误检问题与漏检问题同样突出。天空亮斑、建筑边缘、地面高反差纹理、飞鸟以及远处悬浮物都可能形成与真实飞行目标相似的局部响应。空间上下文建模、区域筛选和注意力增强等代表性方法[6][7][9][10]虽然从不同角度提高了检测精度，但这些方法真正解决的并不只是“看得更清楚”，还包括“把容易混淆的背景排除出去”。对拦截任务而言，误检往往比一般离线识别更具破坏性，因为一次错误观测就可能把后续跟踪与控制链路引向完全错误的方向。',
        '因此，低空小目标检测的可用性不能只依赖单一阈值或单帧分类置信度来判断。ByteTrack[25]的经验说明，低置信度观测并不必然等于无效观测；在某些情况下，它们可能恰恰对应远距离、遮挡或姿态突变下的真实目标。同样的思路也提示检测研究不宜简单追求“高阈值、低误检”的静态平衡，而应结合后端跟踪与状态估计模块共同判断观测是否值得保留。真正有价值的检测输出，不只是某一帧的高分框，而是能够在时间上持续支撑轨迹更新的观测序列。',
        '从这一角度看，检测结果的质量至少包含三个层面：一是目标是否被检出，二是框的位置和尺度是否稳定，三是输出形式是否便于后续模块利用。HOTA[31]对定位与关联质量的综合刻画，实际上也在侧面提醒研究者，感知模块的价值最终要通过连续任务来体现。对于低慢小目标研究而言，未来的检测方法需要把误检抑制、置信度校准和时序可利用性放到同等重要的位置，而不能只围绕静态 AP 展开优化。',
    ])
    heading(doc, '2.5 训练策略、数据增强与先验利用', level=2)
    add_paragraphs(doc, [
        '除网络结构本身外，训练策略对低空小目标检测的影响同样不可忽视。Focal Loss[2]从类别不平衡角度改善了正负样本学习，EfficientDet[3]强调尺度分配与融合效率，QueryDet[5]则把稀疏查询思想引入高分辨率推理过程。这些工作共同说明，小目标性能并不完全由主干网络决定，样本分配方式、候选区域覆盖策略和训练阶段的难样本组织，往往会直接影响目标能否在极少像素条件下被稳定识别。',
        '在无人机和低空场景中，数据增强策略的作用尤其突出。随机裁剪、尺度抖动、背景扰动、亮度变化和模糊模拟等操作，能够在一定程度上逼近真实飞行环境中的距离变化、姿态变化和成像退化。与此同时，增强策略也需要保持克制：若对极小目标进行过强裁剪或缩放，原本有限的目标信息可能被进一步破坏，反而削弱模型对真实目标的辨识能力。因此，小目标任务中的数据增强更适合围绕目标保持、背景多样化和弱纹理保真展开，而不是简单套用通用检测任务中的重增强套路。',
        '另外，任务先验的使用也值得重视。低空飞行目标通常具有特定运动边界、出现区域和尺度变化规律，这些信息既可以体现在数据采样中，也可以体现在候选筛选与后处理环节中。对工程系统来说，合理利用先验并不意味着削弱模型泛化能力，反而往往能够降低误检率并提高系统整体可控性。检测模型若能与任务先验形成互补，其实际价值通常会高于单纯追求通用排行榜成绩的方法。',
    ])
    heading(doc, '3 数据集与评测体系')
    heading(doc, '3.1 无人机视觉 benchmark 的建立与作用', level=2)
    add_paragraphs(doc, [
        '如果说检测器论文主要回答“如何看见目标”，那么 benchmark 研究回答的则是“如何判断一种方法是否真正有效”。UAVDT[13]和 Detection and Tracking Meet Drones Challenge[12]是该方向的重要里程碑。它们把目标检测、单目标跟踪和多目标跟踪统一到无人机视角下进行评测，使研究者开始系统认识到：同样的算法在地面监控场景与无人机视角下可能表现完全不同。俯视视角、目标尺度骤减和平台自身运动共同改变了问题本身。',
        'VisDrone 系列挑战结果[21][22]进一步推动了研究社区对无人机图像检测问题的共识形成。挑战赛虽然更多呈现排行榜和方案分布，但其意义并不只是“谁的指标更高”，更在于它让研究者能够在统一规则下观察各种方法的真实优势与局限。通过这些挑战结果可以看到，多尺度增强、强主干网络和精细训练策略逐步成为基础配置，单一的小修小补很难再产生决定性优势。',
        '基准的价值还体现在问题边界的塑造上。一个领域会优先解决什么问题，很大程度上取决于 benchmark 将什么定义为“成功”。在 UAVDT 和 VisDrone 体系中，检测精度、跟踪连续性和视频理解成为核心指标，因此研究逐渐围绕这些任务展开。换言之，benchmark 不只是被动承载算法比较，而是主动决定研究方向。很多后来被视为标准配置的训练技巧和模型设计，最初正是在统一 benchmark 的反复对比中逐步固化下来的。',
    ])

    heading(doc, '3.2 反无人机与特种场景数据集的出现', level=2)
    add_paragraphs(doc, [
        '随着研究重心从“无人机视角看地面目标”转向“视觉系统看空中飞行目标”，反无人机 benchmark 的重要性显著上升。Vision-Based Anti-UAV Detection and Tracking[14]与 Anti-UAV[15]将低空飞行器本身作为被观测对象，把研究真正推进到低慢小空中目标的层面。与传统无人机视觉数据不同，这类数据中的目标更小、背景更复杂、出入视场更频繁，也更接近拦截和防御场景中的真实需求。',
        'A Unified Transformer-Based Tracker for Anti-UAV Tracking[16]、WebUAV-3M[19]等工作则进一步表明，反无人机任务不仅是一个检测问题，更是一个时序建模和长期跟踪问题。空中飞行器在长时间序列中容易发生尺度突变、背景切换和短时遮挡，因此仅依赖单帧检测很难满足任务需求。相关数据集正是通过长序列、复杂背景和更严格任务定义，把这种需求显性化。',
        '此外，SeaDronesSee[17]与 BIRDSAI[18]展示了特种场景数据集对研究方向的拓展作用。前者面向海上目标，背景单调却易受反光、波纹和距离变化影响；后者引入热红外模态，使低照度和弱纹理目标感知成为可能。这些数据集虽然并不直接等同于低慢小飞行器拦截场景，但它们共同说明：一旦任务环境发生变化，视觉系统所面临的问题也会随之改变，数据集必须承担反映这些变化的角色。',
    ])

    heading(doc, '3.3 现有评测体系的不足', level=2)
    add_paragraphs(doc, [
        '尽管 benchmark 体系已经相当丰富，但从闭环任务视角看，现有评测仍存在明显不足。首先，大多数 benchmark 仍以检测 AP、跟踪成功率或 MOT 指标为中心，而这些指标并不总能反映系统对控制任务的支持程度。一个检测器可能在 AP 上表现良好，却由于时序波动严重而不适合直接用于控制。',
        '其次，当前 benchmark 较少显式考虑目标失锁恢复、连续逼近、图像边缘保持和控制延迟等问题。对于以拦截为目的的研究来说，这些现象往往比单帧定位误差更关键，因为它们直接决定系统是否能够稳定持续地接近目标。',
        '再次，现有数据集虽然覆盖了可见光、热红外、海上和一般低空等多种环境，但仍较少把感知、预测与控制过程放在统一框架下评价。这意味着研究者在感知模块上获得的性能提升，并不一定能够直接转化为系统级收益。也正因如此，后续研究有必要发展更贴近任务需求的综合评测体系。',
    ])


    heading(doc, '3.4 数据偏差、标注粒度与跨场景泛化', level=2)
    add_paragraphs(doc, [
        '不同 benchmark 之间的差异，远不只是样本数量或图像分辨率的不同。UAVDT[13]、VisDrone[21][22]更强调复杂城市场景中的检测与跟踪，Anti-UAV[15]与相关反无人机数据集[14][16]则更加突出空中目标本身的细粒度辨识，SeaDronesSee[17]和 BIRDSAI[18]又分别引入海面反光与热红外成像条件。这意味着同样被统称为“低空小目标”的任务，在数据分布上其实存在明显差异。模型若只在单一数据集上取得较好结果，并不能自动说明其具备跨场景适用性。',
        '标注粒度同样会深刻影响模型训练与评测结果。对尺寸极小的目标而言，边界框只要偏移几个像素，交并比就可能发生明显变化；当目标本身只有有限像素面积时，标注者对边界范围的理解差异甚至会直接改变样本难度。也正因如此，小目标任务中的评测波动往往比一般目标检测更显著。研究者如果忽略标注噪声、目标定义差异和不同数据集之间的类别映射问题，就很容易把数据偏差误判为算法能力差异。',
        '跨场景泛化问题在低慢小目标研究中尤为突出。可见光模型未必适用于热红外数据，海面背景下有效的纹理线索也未必能迁移到城市或郊野环境，针对无人机俯视地面目标训练出来的特征更不一定适合地面设备反向观测空中飞行器。由此可见，未来评价模型优劣时，不能只看单一 benchmark 上的局部领先，而应更多考察其在多场景、多模态和不同成像条件下的稳定性。只有把数据偏差和泛化边界说清楚，相关结论才更具解释力。',
    ])
    heading(doc, '4 目标跟踪与状态估计')
    heading(doc, '4.1 检测后关联路线的发展', level=2)
    add_paragraphs(doc, [
        '在时序处理层面，检测后关联是最经典也最具工程实用性的路线。Deep SORT[24]在传统 SORT 的运动预测框架上引入外观特征，使在线多目标跟踪从单纯依赖几何信息走向几何与外观联合建模。这一改进的意义不只是提升 ID 保持能力，更在于证明了在线跟踪可以在较轻量计算成本下获得可用性能，因此被广泛作为后续研究和工程系统的基线。',
        'ByteTrack[25]的重要贡献在于重新解释低置信度检测框的作用。以往系统往往在检测阈值阶段直接过滤低分框，但 ByteTrack 证明，很多低分框恰恰对应被遮挡、远距离或暂时模糊的真实目标。如果过早删除这些观测，轨迹就会频繁中断。对低空小目标而言，这一思想尤其有价值，因为小目标在远距离和复杂背景下本来就容易产生不稳定置信度。',
        'OC-SORT[30]进一步从观测更新角度增强 SORT 体系的鲁棒性，说明检测后关联路线并没有因为深度端到端方法的兴起而失去价值。相反，随着时序任务变得更复杂，如何在保持轻量和可解释的同时提高观测利用效率，反而变得更加重要。',
        '这一类方法的共同特点在于模块清晰、调试方便、与检测器耦合自然，因此很适合需要与控制系统联动的场景。但它们的局限也很明确：一旦检测器长期不稳定，后续关联质量仍会迅速下降。因此，检测后关联路线的上限在很大程度上取决于前端检测器的可靠性。',
    ])

    heading(doc, '4.2 单目标跟踪与端到端时序建模', level=2)
    add_paragraphs(doc, [
        '与检测后关联不同，另一条研究路线试图把时序建模更深地融入跟踪框架内部。Tracking Objects as Points[26]通过关键点表示统一检测与跟踪，使时序关联不再只是后处理步骤，而是检测过程的一部分。Improving Multiple Object Tracking with Single Object Tracking[29]则借鉴单目标跟踪思想来增强多目标轨迹连续性，说明 MOT 与 SOT 之间并不存在绝对边界。',
        '在单目标跟踪领域，SiamRPN++[27]和 DiMP[28]分别代表了两种典型方向。前者延续了孪生匹配框架，通过更深特征与更合理的定位机制显著提升性能，适合对实时性要求较高的场景；后者把判别式模型预测引入跟踪器更新，在复杂背景和外观变化条件下具有更强鲁棒性。二者虽然主要服务于单目标任务，但其关于时序连续性和局部目标建模的思想，对低空飞行目标持续跟踪同样具有启示意义。',
        '从方法演进来看，端到端时序建模的优势在于能够更直接地利用跨帧信息，减少检测与跟踪模块之间的信息损失；但代价通常是更高的模型复杂度和更大的推理成本。对于机载或实时控制系统而言，这种代价并不总是可接受。因此，不同方法并不存在绝对优劣，而是取决于任务对实时性、鲁棒性和系统复杂度的具体要求。',
    ])

    heading(doc, '4.3 状态估计与跟踪评价', level=2)
    add_paragraphs(doc, [
        '除具体跟踪框架外，状态估计和评价指标也是理解该方向不可忽视的部分。HOTA[31]通过同时考虑检测、关联和定位质量，弥补了单一 MOTA/IDF1 指标对轨迹质量刻画不足的问题。它提示研究者，跟踪性能并不是单一数字可以完全概括的，尤其是在小目标场景中，定位误差、轨迹中断和关联错误往往同时存在。',
        '对于低空小目标任务来说，状态估计的重要性甚至高于一般视频跟踪问题。原因在于控制器并不直接使用“目标是否被检测到”这一二值信息，而更依赖目标位置、速度趋势和短时运动方向是否连续可用。一旦检测结果出现抖动或短时缺失，系统就需要依靠运动模型维持状态连续性。尽管卡尔曼滤波并非最新颖的方法，但大量研究都以不同方式说明，轻量、稳定、可解释的状态预测依然是工程系统中不可或缺的一环。',
        '从这一意义上说，跟踪研究的真正价值并不只是生成更漂亮的轨迹，而在于把不稳定的视觉观测转化为足以支撑后续控制决策的连续状态信息。对于低慢小目标拦截而言，这恰恰是连接感知与控制的关键桥梁。只有当目标状态在时间上可连续解释，后续控制器才可能在观测缺失、目标尺度突变或背景干扰增强时保持平稳响应。',
    ])


    heading(doc, '4.4 失锁恢复、短时预测与轨迹连续性', level=2)
    add_paragraphs(doc, [
        '低空飞行目标在真实视频序列中很少以理想方式持续出现。目标可能因为急转、加减速、云层遮挡、背景高亮区域干扰或摄像机快速机动而短时消失，随后又重新进入可观测状态。因此，跟踪问题并不只是“如何持续关联已经检测到的目标”，还包括“目标短时不可见后如何重新建立稳定轨迹”。在这一点上，Deep SORT[24]、ByteTrack[25]、OC-SORT[30]等方法关于观测利用与轨迹更新的思路，与 SiamRPN++[27]、DiMP[28]等单目标跟踪器关于局部目标重定位的经验，实际上具有很强互补性。',
        '失锁恢复之所以重要，是因为低慢小目标在图像中常常缺乏足够纹理，重新出现时其外观又可能发生明显变化。如果系统过度依赖瞬时外观相似性，就容易在背景相似区域发生误关联；如果系统过度依赖运动模型，又可能在目标突发机动时迅速偏离真实轨迹。现有研究的启示在于，轨迹连续性需要由多种信息共同维持，包括短时运动趋势、历史外观表征、当前观测置信度以及场景约束等。单一线索通常不足以支撑复杂低空场景下的稳定恢复。',
        '对于面向拦截的任务而言，短时预测的重要性还要进一步提高。控制器并不只关心目标当前位置，更关心下一个控制周期内目标可能出现在什么位置。如果预测时域过短，系统会不断被观测延迟牵着走；如果预测过于激进，又容易在目标机动变化时引入额外误差。因此，状态估计模块的价值不仅在于“补洞”，更在于为控制链路提供平滑、连续、可提前量化的目标运动信息。跟踪系统是否具备可靠的短时预测能力，往往直接决定了后续控制是否稳定。',
    ])
    heading(doc, '5 视觉伺服、目标追踪与拦截控制')
    heading(doc, '5.1 图像误差闭环与可见性约束', level=2)
    add_paragraphs(doc, [
        '视觉伺服研究关注的核心问题，是如何把图像空间中的目标偏差转化为飞行器的控制输入。Position-Based Visual Servoing[35]表明，仅知道目标大致位置并不足以形成稳定控制，还需要考虑位姿关系、飞行器动力学和目标运动对闭环系统的影响。Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs[32]则进一步指出，目标可见性并不是检测阶段的问题，而是控制设计本身必须显式考虑的约束。',
        'Image-Based Visual Servoing of Rotorcrafts to Planar Visual Targets of Arbitrary Orientation[34]和 Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets[33]把这一问题推广到更复杂的目标姿态和飞行状态，说明图像误差并不能简单等同于控制误差。目标在图像中的位置变化既受飞行器运动影响，也受目标姿态、深度变化和相机投影关系影响，因此有效控制必须建立在更完整的视觉几何分析之上。',
        '这组研究的共同启示在于，感知与控制之间存在明显语义差距。检测器输出的是边界框、关键点或图像坐标，而控制器需要的是速度指令、姿态调整和稳定性保证。视觉伺服研究的价值，正是在于构建这两种信息之间的映射关系。',
    ])

    heading(doc, '5.2 从视觉追踪到自主拦截', level=2)
    add_paragraphs(doc, [
        '如果说视觉伺服主要解决“如何持续面向目标”，那么自主追踪与拦截研究更进一步关注“如何接近并处置目标”。Autonomous Target Tracking of UAV Using High-Speed Visual Feedback[37]强调高频视觉反馈在动态目标追踪中的意义，说明当控制回路与感知回路耦合后，刷新率本身就是影响系统性能的重要因素。视觉系统不仅要足够准确，还要足够快。',
        'Autonomous Drone Hunter Operating by Deep Learning and All-Onboard Computations in GPS-Denied Environments[36]则从更系统的角度展示了机载自主追踪与处置的可能性。在 GPS 受限环境中，仅依赖视觉和机载计算完成目标持续追踪，说明感知、估计与控制已经可以在较高程度上整合到单个平台中。对低慢小目标研究而言，这类工作代表了从“识别目标”向“处理目标”的方向跃迁。',
        'UAV Guidance[39]中的双目引导拦截思路又进一步表明，一旦任务目标从持续跟踪转向终端拦截，对目标深度、相对位姿和接近路径的估计精度要求会显著提高。普通跟踪任务允许目标偶尔偏离画面中心，甚至短时失锁后恢复；而在拦截任务中，这样的误差往往直接影响终端动作是否成功。',
        'Robust Visual Servoing Formation Tracking Control for Quadrotor UAV Team[38]虽然研究对象是编队控制，但它说明了另一种值得注意的趋势，即视觉控制并不局限于单机单目标，还可能扩展到多机协同、队形保持和联合处置。随着任务复杂度增加，视觉感知系统需要承担的稳定性责任也会进一步提高。',
    ])

    heading(doc, '5.3 仿真验证的意义与边界', level=2)
    add_paragraphs(doc, [
        'AirSim[40]及类似高保真平台在这一研究方向中具有特殊地位。它们既能提供丰富的场景建模能力，也能为图像采集、控制验证和参数调试提供可重复环境。对于感知与控制高度耦合的问题来说，仿真不仅节省实验成本，更重要的是可以帮助研究者在较低风险下快速验证不同策略的影响。',
        '但仿真平台的价值不应被夸大。纹理细节、天气扰动、传感噪声以及真实环境中的突发因素，都可能使仿真表现和真实部署之间存在差距。因此，仿真更适合被视为闭环研究中的一个中间层：它既比纯离线数据实验更接近系统应用，又不能完全取代真实环境验证。把仿真结果与数据集评测、实际飞行试验和系统分析结合起来，才更符合这一方向的发展需求。对于系统级研究而言，仿真平台的意义尤其在于缩短感知算法、状态估计与控制策略联调的迭代周期，而不是简单替代真实测试。',
    ])


    heading(doc, '5.4 感知-控制耦合中的时延与稳定性问题', level=2)
    add_paragraphs(doc, [
        '一旦感知结果被直接送入控制回路，检测和跟踪中的时延问题就不再只是工程实现细节，而会转化为闭环稳定性问题。图像采集、神经网络推理、目标关联、状态估计和控制计算之间的每一个延迟环节，都会累积成控制系统中的相位滞后。Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs[32]、Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets[33]以及高频视觉反馈追踪研究[37]都从不同角度说明，视觉控制的关键并不只在于“能否看见目标”，还在于“能否及时、连续地把观测转化为稳定控制量”。',
        '这种耦合关系也解释了为什么一些离线指标优异的方法进入闭环系统后并不一定表现最好。检测框如果出现轻微抖动，在静态评测中可能几乎不影响 AP；但在控制回路中，它可能被持续放大为姿态微振和轨迹摆动。相反，一些绝对精度并非最高但输出更平滑、延迟更低的方法，反而更适合作为控制前端。这说明感知算法的优劣不能脱离其使用方式来评价，特别是在视觉伺服和拦截场景中，输出平稳性往往与检测准确性同样重要。',
        '此外，终端逼近阶段还会引入更严格的安全边界。UAV Guidance[39]所体现的拦截过程并不只是简单追上目标，而是要求系统在保持目标可见、估计相对位置并控制接近速度的同时，避免由于观测抖动造成过度机动。Autonomous Drone Hunter[36]和编队视觉控制研究[38]进一步表明，当任务升级为自主处置或多机协同时，感知链路中的不稳定性会被进一步放大。由此可见，时延补偿、控制饱和约束和安全保持策略，都是感知控制耦合研究中无法回避的核心问题。',
    ])
    heading(doc, '6 综合讨论')
    heading(doc, '6.1 研究主线已经从单点优化转向模块协同', level=2)
    add_paragraphs(doc, [
        '低空小目标研究最明显的变化，是问题表述方式从单点算法优化逐步走向跨模块协同。早期工作更多关注检测器如何提升对小目标的识别能力，而近年的研究则越来越重视“检测结果如何被跟踪利用”“跟踪状态如何支撑控制决策”“控制过程如何反过来影响目标可见性”等系统性问题。',
        '这种变化意味着研究目标已经从“某一个模块指标更高”转向“整条技术链路更稳定”。对拦截任务尤其如此，因为拦截的成功并不只取决于某一帧是否检测到目标，而取决于系统能否在连续时间内稳定感知目标、保持状态估计并输出合理控制。',
        '也正因此，检测、跟踪、状态估计和控制几个方向不能孤立理解。只有把它们放到统一任务背景下分析，才能理解各自的价值与边界。',
    ])

    heading(doc, '6.2 当前研究仍存在的主要问题', level=2)
    add_paragraphs(doc, [
        '第一，极小目标与复杂背景之间的矛盾尚未真正解决。多尺度特征融合、注意力增强和上下文建模虽已明显改善检测能力，但在远距离、逆光、背景纹理相似和运动模糊叠加的情况下，漏检和误检依然频繁出现。对于低慢小目标任务而言，这一问题并不是局部性能波动，而是系统风险来源。',
        '第二，时序稳定性依旧是制约系统可用性的核心因素。许多检测器在离线评测中指标较好，但一旦进入连续视频流，其输出会出现中心漂移、尺度抖动或短时消失。跟踪和状态估计虽然能够提供补偿，但如果前端视觉观测长期不稳定，后端模块的修正能力也会迅速耗尽。',
        '第三，现有 benchmark 对控制层需求覆盖不足。大量研究围绕检测 AP、跟踪成功率或 MOT 指标展开，但对“是否足以支撑持续逼近”“是否能够在目标短时失锁后迅速恢复”“是否会引起控制振荡”等问题关注较少。对于面向拦截的应用来说，这些问题比单纯的离线识别精度更关键。',
        '第四，仿真到真实环境之间仍存在明显鸿沟。很多方法在数据集和仿真环境中表现稳定，但在真实空域中会受到风扰动、传感噪声、镜头模糊和光照变化影响，从而出现性能退化。如何通过域随机化、合成数据与真实数据混合训练、多模态融合等方法缩小这一差距，仍是重要课题。',
    ])

    heading(doc, '6.3 对低慢小目标拦截研究的启示', level=2)
    add_paragraphs(doc, [
        '面向低慢小目标拦截的研究不宜把注意力仅仅放在检测模型本身。真正具有系统意义的研究路径，应当在检测、短时预测和控制闭环之间建立更紧密联系。检测器需要尽量输出对控制更友好的观测量，状态预测模块需要在漏检条件下维持连续性，控制器则需要显式考虑目标可见性和图像误差演化。',
        '换句话说，低慢小目标拦截研究既要吸收小目标检测和无人机 tracking 的最新进展，又不能完全照搬那些只服务于离线识别或排行榜竞赛的技术路线。更合理的方向是围绕闭环系统的稳定性、实时性和可复现性进行场景化设计，使研究成果既能够回应这一领域的共性问题，也能为实际系统实现提供更直接的技术依据。',
    ])


    heading(doc, '6.4 方法选择中的若干取舍', level=2)
    add_paragraphs(doc, [
        '通读这一方向的代表性工作可以发现，并不存在一种对所有场景都占优的“万能方法”。多尺度增强检测器在极小目标和复杂背景下具有优势，但往往伴随更高推理成本；检测后关联框架结构清晰、部署方便，却更依赖前端观测质量；端到端时序模型能够更充分利用跨帧信息，但在机载平台上未必总能满足实时性要求。方法选择因此从来不是单纯比较谁的指标更高，而是比较谁更适合当前任务约束。',
        '同样值得注意的是，算法论文中的领先指标并不一定等同于系统中的真实收益。对于低慢小目标任务来说，少量误检、中心点抖动、尺度估计不稳定和偶发失锁，都可能在闭环链路中被逐步放大。也就是说，系统最敏感的往往不是排行榜上最醒目的那一列数字，而是那些在标准评测中不够显眼、却会持续影响后端模块的误差形态。只有把这些误差放回真实使用场景中理解，才能真正判断一种方法的适用范围。',
        '从研究组织方式看，模块化设计与更紧耦合的联合设计仍将长期并存。模块化路线便于调试、定位问题和逐步替换组件，尤其适合工程实现和原型验证；更紧耦合的设计则有机会减少信息损失，在系统层面取得更好性能。对低空小目标拦截而言，较为稳妥的思路往往不是一开始就追求全链路端到端，而是先把检测、跟踪和控制之间的接口定义清楚，再逐步提高模块之间的信息共享程度。这样的研究路径通常更具可解释性，也更容易形成可靠结果。',
    ])
    heading(doc, '6.5 从模块性能到系统容错', level=2)
    add_paragraphs(doc, [
        '对低慢小目标任务而言，系统容错能力往往比某一模块的峰值性能更关键。检测器偶发误检、跟踪器短时漂移、状态预测出现轻微偏差，在离线评测里可能只是局部误差；但一旦这些误差在闭环系统中连续传递，就可能逐步演化为视轴偏移、控制振荡甚至任务失败。因此，评价系统是否成熟，不能只看单个模块在理想条件下的最好结果，还要看整条链路在非理想条件下能否稳定退化。',
        '这也解释了为什么冗余设计在相关系统中具有现实意义。前端检测结果可以为跟踪器提供初始化与重捕获依据，状态预测可以在短时失锁时维持控制连续性，控制器则可以通过限幅、缓启动或保持策略避免对瞬时异常观测过度响应。这样的容错设计虽然不一定显著提升某项单独指标，却能明显提升系统在复杂环境中的生存能力。对于拦截任务来说，这种能力往往比局部精度提升更具实际价值。',
        '从研究方法上看，系统容错意识还要求研究者更关注误差是如何传播的。哪些误差会被后端自然吸收，哪些误差会被逐级放大，哪些情况下系统应主动降级而非继续追踪，这些问题都值得被纳入论文分析框架。只有把误差传递机制讲清楚，相关工作才能从“方法有效”进一步走向“系统可靠”。',
    ])
    heading(doc, '7 发展趋势与展望')
    add_paragraphs(doc, [
        '未来的一个关键方向，是进一步推动感知、预测与控制的协同设计。当前许多工作仍以模块化方式解决问题，这在工程上具有可实现性，但也容易造成信息在模块之间逐层损失。若能够在模型设计阶段就考虑哪些观测量最有利于后续控制、哪些时序特征最值得保留，就可能在系统层面获得更大收益。',
        '第二个方向是面向低算力平台的轻量化实现。无论是小目标检测还是 Transformer 跟踪器，越来越复杂的模型结构都在不断抬高部署门槛。对于真实无人系统而言，模型必须在有限机载资源下稳定运行，因此轻量主干网络、稀疏推理、模型压缩和软硬件协同优化将长期受到重视。',
        '第三个方向是更贴近任务需求的 benchmark 构建。未来的评测体系不应只停留在检测或跟踪指标层面，而应考虑图像可见性保持、短时失锁恢复、控制延迟敏感性以及系统级成功率等更具任务意义的指标。只有这样，benchmark 才能真正引导研究向闭环系统发展。',
        '第四个方向是仿真到真实环境的迁移增强。仿真平台将继续在算法开发和系统联调中发挥重要作用，但如何通过数据增强、域随机化、风场与光照建模以及有限真实数据微调，减少仿真与实装之间的性能落差，仍会是低空小目标研究绕不开的问题。',
        '总体而言，低空小目标研究不会停留在单一的视觉识别层面，而会继续向更加系统化、任务化和工程化的方向演进。谁能更好地把检测、状态估计、控制和任务约束统一起来，谁就更可能在这一方向取得具有实际价值的突破。',
    ])


    heading(doc, '7.1 多源感知融合与弱目标增强', level=2)
    add_paragraphs(doc, [
        '从现有研究基础看，单一可见光模态虽然仍是主流，但其局限也越来越明显。逆光、低照度、云层遮挡、复杂背景以及远距离成像条件都会显著削弱可见光特征的稳定性。SeaDronesSee[17]和 BIRDSAI[18]所体现出的海面反光、热红外对比和弱纹理成像差异，说明不同模态对低慢小目标具有不同敏感性。未来若希望进一步提高弱目标可感知性，多源信息融合几乎是不可回避的方向。',
        '但多源融合并不意味着简单叠加传感器数量。模态之间的时间同步、空间标定、分辨率差异和噪声特性都会给系统带来新的工程负担。如果融合策略过于复杂，机载实时性又会受到明显影响。因此，更有价值的研究应当围绕“哪些信息对任务最关键”来设计融合方式，例如在低照度条件下突出热特征，在复杂背景中强化运动一致性，在远距离条件下利用历史时序信息弥补单帧表观不足。',
        '从任务角度看，多源融合的意义最终仍在于提高连续可用性，而不仅仅是提升某项离线指标。只有当多模态信息能够帮助系统更稳定地检出目标、减少误检并改善失锁后的恢复能力时，这类方法才真正具有系统价值。对低慢小目标拦截而言，多源感知的研究重点应当是围绕弱目标增强和闭环鲁棒性展开，而不是追求与任务脱节的复杂结构堆叠。',
    ])

    heading(doc, '7.2 从离线指标走向闭环任务评价', level=2)
    add_paragraphs(doc, [
        '未来研究的另一关键方向，是建立更符合系统需求的评价方法。当前大量工作仍停留在 AP、成功率、MOTA 或 IDF1 等传统指标层面，这些指标固然重要，却不足以完整反映感知结果在控制链路中的真实作用。对于拦截任务而言，更值得关注的问题包括：目标短时失锁后能否快速恢复、图像中心偏差是否持续平稳、控制时延对逼近过程影响多大，以及终端接近阶段是否会因为感知抖动引发不必要机动。',
        'HOTA[31]已经从检测、关联和定位三个维度给出了更综合的评价思路，但面向闭环系统的评价仍有进一步拓展空间。结合 UAVDT[13]、Anti-UAV[15] 等 benchmark 的经验，未来完全可以把数据集评测、视频连续性分析和闭环控制结果联系起来，形成多层次评价框架。这样一来，研究者不仅能够知道某种方法在离线数据上是否有效，也能够更清楚地判断它在真实任务中是否真正可用。',
        '从更长远的角度看，任务导向评价体系还将直接影响研究方向本身。只要 benchmark 仍主要奖励静态检测精度，研究重心就会继续向单点视觉优化聚集；而一旦评价体系把连续跟踪、状态可解释性、时延敏感性和系统成功率纳入核心指标，相关研究自然会更快转向感知、预测与控制的协同设计。评价标准的变化，往往比个别算法更新更能决定一个领域的演化方向。',
    ])

    heading(doc, '7.3 面向真实部署的系统协同', level=2)
    add_paragraphs(doc, [
        '面向真实部署的研究还需要更加重视系统协同问题。Autonomous Drone Hunter[36]和 AirSim[40]所代表的两类工作分别说明了一个事实：一方面，机载自主系统必须在有限算力、有限能耗和有限通信条件下稳定运行；另一方面，仿真平台虽然能够帮助快速迭代，却无法自动替代真实空域中的部署验证。因此，真正具有落地价值的研究，通常不是某一个模块在理想条件下做到极致，而是整条链路在复杂约束下仍然保持可解释、可调试和可复现。',
        '这也意味着未来的系统设计需要更清楚地处理接口问题。检测模块输出什么形式的信息最便于跟踪与控制利用，状态估计应当保留哪些不确定性描述，控制器在感知波动时应采取怎样的安全退化策略，这些问题都值得在系统级研究中被前置考虑。只有把接口设计、资源约束与任务目标一并纳入方法选择，低慢小目标研究才能逐步从“能做演示”走向“可长期稳定运行”。',
        '总体来看，真实部署导向并不是削弱学术研究的深度，恰恰相反，它要求研究者把模型能力、数据条件、硬件资源、任务评价和安全约束放到同一框架中统筹考虑。对低空小目标视觉研究而言，谁能够更早完成这种系统化转变，谁就更可能在后续工作中提出既有学术价值又有工程意义的方法。',
    ])
    heading(doc, '7.4 安全性约束与任务边界', level=2)
    add_paragraphs(doc, [
        '随着研究逐步从检测和跟踪走向追踪与拦截，安全性约束的重要性会越来越突出。低慢小目标任务并不是单纯的视觉识别竞赛，而是与目标判别、动作决策和空间约束紧密相连的系统问题。一旦目标类别判断不准、相对位置估计不稳或控制动作过于激进，就可能在真实应用中引入额外风险。因此，后续研究不能只关注“能否发现目标”，还需要同时关注“是否能够在约束条件下安全处理目标”。',
        '这类约束首先体现为控制边界。飞行器的机动能力、视场范围、最小安全距离、相机盲区以及环境中的障碍物都会限制系统能够采取的动作范围。即便感知结果本身准确，如果控制策略忽视这些边界，系统仍可能在终端阶段出现过度逼近、视场丢失或姿态超调等问题。也正因如此，未来的相关研究更适合把安全约束视作系统设计的内生部分，而不是实验完成后的附加说明。',
        '安全性问题还体现为任务边界的清晰定义。不同场景下，系统究竟是进行预警、持续跟踪、伴飞监视还是终端拦截，其感知精度要求、时延容忍度和控制策略都并不相同。如果任务边界定义含混，研究者就容易在方法评价时混用不同目标，导致结论看似全面却缺乏针对性。对低空小目标视觉研究而言，把任务边界、控制权限和安全退化机制一并说明清楚，将有助于该方向形成更稳健的研究范式。',
    ])
    heading(doc, '8 结论')
    add_paragraphs(doc, [
        '低空小目标视觉研究已经从早期的通用检测问题逐步发展为一个包含数据集建设、时序跟踪、视觉伺服和拦截控制的系统性研究领域。已有研究表明，多尺度建模、上下文增强和注意力机制显著提高了小目标检测能力，benchmark 的持续完善推动了研究标准化，跟踪与状态估计方法加强了时间连续性，而视觉伺服与自主追踪研究则将感知成果进一步引入控制回路。',
        '与此同时，低空小目标任务仍受到漏检、误检、时序抖动、系统延迟和仿真迁移等问题制约。对这一方向的后续研究而言，最值得关注的并不是单点模块的短期性能提升，而是如何构建一条稳定、实时且可复现的感知 - 预测 - 控制闭环链路。围绕这一目标展开的研究，将更有可能真正支撑低慢小目标拦截等实际任务需求。',
    ])

    heading(doc, '参考文献')
    for idx in range(1, 41):
        ref_para(doc, idx, build_ref_text(idx))

    doc.save(OUT_PATH)
    print(OUT_PATH)

if __name__ == '__main__':
    build()






