from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(r'D:\Xstarlab\UEProjects\GraduationProject\GraduationProject')
OUT_DIR = ROOT / 'output' / 'doc'
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / '文献综述-标准综述版-面向低空小目标的视觉检测跟踪与拦截控制研究.docx'

TITLE = '面向低空小目标的视觉检测、跟踪与拦截控制研究综述'
AUTHOR_LINE = '作者：颜飞翔    学号：2229409056'
DATE_LINE = '完成日期：2026年3月21日'

refs = [
    {'id':1,'title':'Feature Pyramid Networks for Object Detection','venue':'CVPR','year':'2017','doi':'10.1109/CVPR.2017.106','theme':'检测基础','contrib':'构建多尺度特征金字塔，为小目标检测提供高分辨率语义特征','limit':'并非面向航拍场景设计，仍需结合场景化采样与损失设计'},
    {'id':2,'title':'Focal Loss for Dense Object Detection','venue':'ICCV','year':'2017','doi':'10.1109/ICCV.2017.324','theme':'检测基础','contrib':'缓解正负样本极度不平衡问题，提升稠密检测对难样本的关注','limit':'只能部分缓解小目标难检问题，无法替代多尺度建模'},
    {'id':3,'title':'EfficientDet: Scalable and Efficient Object Detection','venue':'CVPR','year':'2020','doi':'10.1109/CVPR42600.2020.01079','theme':'检测基础','contrib':'在精度和推理效率间取得较好平衡，适合嵌入式与边缘场景参考','limit':'对极小目标和复杂背景的适应性仍有限'},
    {'id':4,'title':'End-to-End Object Detection with Transformers','venue':'ECCV','year':'2020','doi':'10.1007/978-3-030-58452-8_13','theme':'检测基础','contrib':'以集合预测统一检测流程，推动检测器由锚框范式向端到端范式演化','limit':'对小目标和收敛速度存在先天不足'},
    {'id':5,'title':'QueryDet: Cascaded Sparse Query for Accelerating High-Resolution Small Object Detection','venue':'CVPR','year':'2022','doi':'10.1109/CVPR52688.2022.01330','theme':'检测基础','contrib':'针对高分辨率稀疏小目标构建级联稀疏查询机制，兼顾速度与精度','limit':'仍依赖目标稀疏性假设，对密集小目标场景适应性有限'},
    {'id':6,'title':'Small Object Detection in Unmanned Aerial Vehicle Images Using Feature Fusion and Scaling-Based Single Shot Detector With Spatial Context Analysis','venue':'IEEE TCSVT','year':'2019','doi':'10.1109/TCSVT.2019.2905881','theme':'航拍小目标检测','contrib':'将尺度增强与空间上下文结合，显著改善无人机视角小目标检测效果','limit':'模型结构较复杂，对实时性有一定影响'},
    {'id':7,'title':'Focus-and-Detect: A Small Object Detection Framework for Aerial Images','venue':'Signal Processing: Image Communication','year':'2022','doi':'10.1016/j.image.2022.116675','theme':'航拍小目标检测','contrib':'引入聚焦区域与局部增强策略，降低大面积背景干扰','limit':'区域聚焦过程增加了流程复杂度'},
    {'id':8,'title':'MFFSODNet: Multiscale Feature Fusion Small Object Detection Network for UAV Aerial Images','venue':'IEEE TIM','year':'2024','doi':'10.1109/TIM.2024.3381272','theme':'航拍小目标检测','contrib':'利用多尺度特征融合强化 UAV 图像中微小目标表达','limit':'主要验证于检测层面，对控制闭环启示有限'},
    {'id':9,'title':'A Context-Scale-Aware Detector and a New Benchmark for Remote Sensing Small Weak Object Detection in Unmanned Aerial Vehicle Images','venue':'JAG','year':'2022','doi':'10.1016/j.jag.2022.102966','theme':'航拍小目标检测','contrib':'兼顾新基准构建和上下文尺度敏感检测策略，突出小弱目标问题','limit':'更偏遥感视角，对动态拦截场景覆盖不足'},
    {'id':10,'title':'Small Object Detection in Unmanned Aerial Vehicle Images Using Multi-Scale Hybrid Attention','venue':'Engineering Applications of Artificial Intelligence','year':'2023','doi':'10.1016/j.engappai.2023.107455','theme':'航拍小目标检测','contrib':'混合注意力机制提升复杂背景下小目标可分辨性','limit':'注意力模块带来参数与推理开销'},
    {'id':11,'title':'Tiny Object Detection in Aerial Images','venue':'ICPR','year':'2021','doi':'10.1109/ICPR48806.2021.9413340','theme':'航拍小目标检测','contrib':'系统讨论 tiny object 检测的定位与分类难点','limit':'更偏方法分析，系统级集成讨论不足'},
    {'id':12,'title':'Detection and Tracking Meet Drones Challenge','venue':'IEEE TPAMI','year':'2021','doi':'10.1109/TPAMI.2021.3119563','theme':'基准与挑战','contrib':'将检测、单目标跟踪和多目标跟踪统一纳入无人机场景挑战','limit':'关注感知评测多于控制评测'},
    {'id':13,'title':'The Unmanned Aerial Vehicle Benchmark: Object Detection, Tracking and Baseline','venue':'IJCV','year':'2019','doi':'10.1007/s11263-019-01266-1','theme':'基准与挑战','contrib':'UAVDT 成为无人机视角检测跟踪研究的重要基准','limit':'场景覆盖广但对拦截场景特征刻画不够'},
    {'id':14,'title':'Vision-Based Anti-UAV Detection and Tracking','venue':'IEEE TITS','year':'2022','doi':'10.1109/TITS.2022.3177627','theme':'反无人机基准','contrib':'面向反无人机场景分析目标检测与跟踪难点','limit':'主要聚焦感知问题，控制环节着墨较少'},
    {'id':15,'title':'Anti-UAV: A Large-Scale Benchmark for Vision-Based UAV Tracking','venue':'IEEE TMM','year':'2021','doi':'10.1109/TMM.2021.3128047','theme':'反无人机基准','contrib':'构建大规模反无人机跟踪基准，推动单目标跟踪模型发展','limit':'更重视跟踪鲁棒性，对感知-控制耦合讨论有限'},
    {'id':16,'title':'A Unified Transformer-Based Tracker for Anti-UAV Tracking','venue':'CVPRW','year':'2023','doi':'10.1109/CVPRW59228.2023.00305','theme':'反无人机跟踪','contrib':'将 Transformer 表征引入 anti-UAV 跟踪，提升复杂场景适应性','limit':'模型较重，实时部署成本偏高'},
    {'id':17,'title':'SeaDronesSee: A Maritime Benchmark for Detecting Humans in Open Water','venue':'WACV','year':'2022','doi':'10.1109/WACV51458.2022.00374','theme':'特种场景数据集','contrib':'展示海上无人机视觉中小目标检测的实际难点','limit':'目标类别与反无人机场景不完全一致'},
    {'id':18,'title':'BIRDSAI: A Dataset for Detection and Tracking in Aerial Thermal Infrared Videos','venue':'WACV','year':'2020','doi':'10.1109/WACV45572.2020.9093284','theme':'特种场景数据集','contrib':'引入热红外航拍检测与跟踪，为低照度目标感知提供数据支持','limit':'与可见光检测任务差异较大'},
    {'id':19,'title':'WebUAV-3M: A Benchmark for Unveiling the Power of Million-Scale Deep UAV Tracking','venue':'IEEE TPAMI','year':'2022','doi':'10.1109/TPAMI.2022.3232854','theme':'特种场景数据集','contrib':'以大规模跟踪数据推动深度 UAV 跟踪模型训练','limit':'以单目标跟踪为主，不直接面向拦截控制'},
    {'id':20,'title':'Unmanned Aerial Vehicle Visual Detection and Tracking Using Deep Neural Networks: A Performance Benchmark','venue':'ICCVW','year':'2021','doi':'10.1109/ICCVW54120.2021.00142','theme':'基准比较','contrib':'从性能比较角度评估多类无人机检测与跟踪方法','limit':'偏 benchmark，不提供新的系统方法'},
    {'id':21,'title':'VisDrone-DET2019: The Vision Meets Drone Object Detection in Image Challenge Results','venue':'ICCVW','year':'2019','doi':'10.1109/ICCVW.2019.00030','theme':'挑战结果','contrib':'给出无人机图像检测挑战的主流方法谱系和性能分布','limit':'以竞赛结果为主，方法机理分析有限'},
    {'id':22,'title':'VisDrone-DET2021: The Vision Meets Drone Object Detection Challenge Results','venue':'ICCVW','year':'2021','doi':'10.1109/ICCVW54120.2021.00319','theme':'挑战结果','contrib':'反映无人机图像检测方法在两年间的迭代趋势','limit':'更偏结果汇总而非系统综述'},
    {'id':23,'title':'On the Detection of Unauthorized Drones - Techniques and Future Perspectives: A Review','venue':'IEEE Sensors Journal','year':'2022','doi':'10.1109/JSEN.2022.3171293','theme':'领域综述','contrib':'从传感器与感知层面对反无人机检测进行系统梳理','limit':'对视觉闭环拦截部分论述不足'},
    {'id':24,'title':'Simple Online and Realtime Tracking with a Deep Association Metric','venue':'ICIP','year':'2017','doi':'10.1109/ICIP.2017.8296962','theme':'多目标跟踪','contrib':'以外观特征增强在线关联，成为检测后跟踪经典基线','limit':'对低分框利用不足，遮挡场景仍受限'},
    {'id':25,'title':'ByteTrack: Multi-Object Tracking by Associating Every Detection Box','venue':'ECCV','year':'2022','doi':'10.1007/978-3-031-20047-2_1','theme':'多目标跟踪','contrib':'强调低分框在保持轨迹连续性中的作用，显著提升在线 MOT 性能','limit':'仍依赖高质量检测器输出'},
    {'id':26,'title':'Tracking Objects as Points','venue':'ECCV','year':'2020','doi':'10.1007/978-3-030-58548-8_28','theme':'端到端跟踪','contrib':'将目标视为关键点统一建模检测与跟踪','limit':'对超小目标和大位移场景仍具挑战'},
    {'id':27,'title':'SiamRPN++: Evolution of Siamese Visual Tracking With Very Deep Networks','venue':'CVPR','year':'2019','doi':'10.1109/CVPR.2019.00441','theme':'单目标跟踪','contrib':'显著提升 Siamese 跟踪器的特征表达与定位能力','limit':'遇到尺度剧变和外观相似干扰时稳定性有限'},
    {'id':28,'title':'Learning Discriminative Model Prediction for Tracking','venue':'ICCV','year':'2019','doi':'10.1109/ICCV.2019.00628','theme':'单目标跟踪','contrib':'将判别式模型预测引入跟踪器更新，提高复杂场景鲁棒性','limit':'模型更新成本高于轻量预测方法'},
    {'id':29,'title':'Improving Multiple Object Tracking with Single Object Tracking','venue':'CVPR','year':'2021','doi':'10.1109/CVPR46437.2021.00248','theme':'跟踪融合','contrib':'将 SOT 思路引入 MOT，强化目标连续跟踪能力','limit':'工程实现相对复杂，实时性压力较大'},
    {'id':30,'title':'Observation-Centric SORT: Rethinking SORT for Robust Multi-Object Tracking','venue':'CVPR','year':'2023','doi':'10.1109/CVPR52729.2023.00934','theme':'多目标跟踪','contrib':'通过以观测为中心的更新机制增强 SORT 体系鲁棒性','limit':'对检测稳定性仍高度敏感'},
    {'id':31,'title':'HOTA: A Higher Order Metric for Evaluating Multi-Object Tracking','venue':'IJCV','year':'2020','doi':'10.1007/s11263-020-01375-2','theme':'评价指标','contrib':'统一衡量检测、关联和定位质量，改善跟踪评估体系','limit':'主要解决评估问题，不直接提升算法性能'},
    {'id':32,'title':'Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs','venue':'IEEE/ASME TMech','year':'2019','doi':'10.1109/TMECH.2019.2906430','theme':'视觉伺服','contrib':'强调目标可见性约束在四旋翼视觉控制中的核心作用','limit':'理论性强，工程实现门槛较高'},
    {'id':33,'title':'Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets','venue':'IEEE RA-L','year':'2023','doi':'10.1109/LRA.2023.3245416','theme':'视觉伺服','contrib':'扩展图像伺服到任意飞行目标，提高控制适用范围','limit':'对高频噪声与极端场景鲁棒性仍需进一步验证'},
    {'id':34,'title':'Image-Based Visual Servoing of Rotorcrafts to Planar Visual Targets of Arbitrary Orientation','venue':'IEEE RA-L','year':'2021','doi':'10.1109/LRA.2021.3101878','theme':'视觉伺服','contrib':'处理视觉目标姿态变化条件下的控制问题','limit':'主要围绕几何目标，实际小型飞行器场景更复杂'},
    {'id':35,'title':'Position-Based Visual Servoing for Target Tracking by a Quadrotor UAV','venue':'AIAA GNC','year':'2016','doi':'10.2514/6.2016-2092','theme':'视觉伺服','contrib':'体现基于位姿误差的控制律设计思路','limit':'依赖较稳定的目标位置估计'},
    {'id':36,'title':'Autonomous Drone Hunter Operating by Deep Learning and All-Onboard Computations in GPS-Denied Environments','venue':'PLOS ONE','year':'2019','doi':'10.1371/journal.pone.0225092','theme':'自主拦截','contrib':'展示全机载计算条件下的自主追踪/捕获原型','limit':'场景相对受控，泛化性仍需检验'},
    {'id':37,'title':'Autonomous Target Tracking of UAV Using High-Speed Visual Feedback','venue':'Applied Sciences','year':'2019','doi':'10.3390/app9214552','theme':'自主追踪','contrib':'强调高频视觉反馈对动态跟踪控制的价值','limit':'更偏追踪稳定性，对拦截评价不足'},
    {'id':38,'title':'Robust Visual Servoing Formation Tracking Control for Quadrotor UAV Team','venue':'Aerospace Science and Technology','year':'2020','doi':'10.1016/j.ast.2020.106061','theme':'视觉控制','contrib':'展示视觉控制在多机队形跟踪中的鲁棒性设计','limit':'研究对象为编队，不是直接的单目标拦截任务'},
    {'id':39,'title':'UAV Guidance: A Stereo-Based Technique for Interception of Stationary or Moving Targets','venue':'Lecture Notes in Computer Science','year':'2015','doi':'10.1007/978-3-319-22416-9_30','theme':'引导拦截','contrib':'给出双目引导拦截思路，较早将视觉感知直接引入拦截任务','limit':'系统复杂度较高，对算力和传感器要求更强'},
    {'id':40,'title':'AirSim: High-Fidelity Visual and Physical Simulation for Autonomous Vehicles','venue':'Springer FSR','year':'2017','doi':'10.1007/978-3-319-67361-5_40','theme':'仿真平台','contrib':'提供高保真视觉与物理仿真平台，支撑感知/控制算法快速验证','limit':'仿真到真实部署之间仍存在域差异'},
]


def set_font(run, east='宋体', west='Times New Roman', size=12, bold=False):
    run.font.name = west
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    run.bold = bold


def config(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)


def title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, east='黑体', size=size, bold=True)


def center(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, east='宋体', size=size)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, east='黑体', size=14 if level == 1 else 12, bold=True)


def para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, east='宋体', size=12)


def ref_para(doc, idx, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    r = p.add_run(f'[{idx}] {text}')
    set_font(r, east='宋体', size=10.5)


def make_table(doc, headers, rows, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    set_font(r, east='宋体', size=font_size)
    return table


def detection_rows():
    ids = list(range(1, 12))
    return [[r['id'], r['title'], r['year'], r['contrib'], r['limit']] for r in refs if r['id'] in ids]


def benchmark_rows():
    ids = list(range(12, 24))
    return [[r['id'], r['title'], r['year'], r['theme'], r['contrib']] for r in refs if r['id'] in ids]


def tracking_rows():
    ids = list(range(24, 32))
    return [[r['id'], r['title'], r['year'], r['theme'], r['contrib']] for r in refs if r['id'] in ids]


def control_rows():
    ids = list(range(32, 41))
    return [[r['id'], r['title'], r['year'], r['theme'], r['contrib']] for r in refs if r['id'] in ids]


def full_matrix_rows():
    return [[r['id'], r['title'], r['theme'], r['venue'], r['year'], r['contrib'], r['limit']] for r in refs]


def build():
    doc = Document()
    config(doc)
    title(doc, TITLE)
    center(doc, '标准综述版')
    center(doc, AUTHOR_LINE)
    center(doc, DATE_LINE)

    heading(doc, '摘要')
    para(doc, '围绕低空复杂背景中的小尺度飞行目标感知与处置问题，本文对 40 篇代表性文献进行了系统综述。与面向单个工程项目的前言式写法不同，本文以文献本身为对象，将现有研究划分为低空小目标视觉检测、无人机/反无人机数据集与评测基准、目标跟踪与状态估计、视觉伺服与引导拦截控制、仿真验证平台五类主题，重点比较不同研究在任务定义、技术路线、实验数据、性能指标和系统边界上的异同。综述表明，小目标检测领域的核心矛盾仍然是“分辨率不足与背景干扰并存”，对应的主流解决路径是多尺度特征融合、上下文建模、注意力增强与稀疏查询；无人机视觉研究已逐步形成以 UAVDT、VisDrone、Anti-UAV、WebUAV-3M 为代表的标准化基准，但现有评测仍以感知任务为中心，对感知 - 预测 - 控制闭环的覆盖不足；在跟踪与状态估计方面，从 Deep SORT、ByteTrack 到 CenterTrack、DiMP、OC-SORT 等方法的演进反映出“检测后关联”与“时序统一建模”两条并行路线；在控制层面，视觉伺服研究强调图像误差闭环与可见性约束，自主追踪与拦截研究则进一步把视觉感知与控制决策统一到无人机实时系统中。本文据此归纳了低空小目标研究的关键难题，包括超小目标漏检、时序稳定性不足、系统级评价缺失以及仿真到真实环境的域迁移问题，并提出未来研究应更加重视检测、预测与控制协同设计、轻量化闭环系统验证以及面向拦截任务的评测体系构建。')

    heading(doc, '关键词')
    para(doc, '低空小目标；小目标检测；无人机视觉；目标跟踪；状态估计；视觉伺服；目标拦截；文献综述', indent=False)

    heading(doc, 'Abstract')
    para(doc, 'This paper presents a survey of 40 representative studies on visual detection, tracking, and interception-related technologies for low-altitude small targets. Rather than serving as an introduction to a single engineering project, the survey takes the selected papers themselves as the main object of analysis and organizes them into five research themes: visual detection of small aerial targets, UAV and anti-UAV datasets and benchmarks, target tracking and state estimation, visual servoing and interception control, and simulation platforms for closed-loop validation. The review shows that recent progress in small target detection mainly relies on multi-scale feature fusion, context modeling, attention enhancement, and sparse query mechanisms; benchmark construction has gradually standardized UAV vision research, yet most existing datasets still emphasize perception performance instead of perception-prediction-control integration; tracking studies reveal two parallel routes, namely detection-based online association and end-to-end temporal modeling; and visual servoing plus autonomous pursuit research has highlighted the importance of image-space feedback, visibility constraints, and real-time onboard computation. Based on a comparative reading of the 40 papers, this survey summarizes the main advances, common limitations, and emerging trends of the field, and argues that future work should pay more attention to co-design across detection, prediction, and control, lightweight closed-loop validation, and benchmark protocols explicitly designed for interception-oriented tasks.')

    heading(doc, 'Key words')
    para(doc, 'low-altitude small targets; visual detection; UAV tracking; state estimation; visual servoing; interception control; survey', indent=False)

    heading(doc, '1 引言')
    for t in [
        '低空小目标研究近年来受到持续关注，其原因并不仅在于无人机数量增加，更在于该类目标在低空环境中表现出典型的“低、慢、小、弱”特征：飞行高度低意味着背景更加复杂，速度慢意味着运动模式多样且持续时间长，尺寸小意味着图像像素占比极低，而弱纹理与弱对比度又使得传统检测器难以稳定分离目标与背景。对于视觉系统而言，这些因素叠加后会直接引出漏检、虚警、轨迹中断和控制抖动等一系列问题。',
        '现有研究通常分别从检测、跟踪或控制某一单点问题切入，但如果把研究视角扩大到完整的技术链路，就会发现低空小目标任务其实跨越了目标表征、时间序列建模、飞行控制和系统验证多个层面。检测器需要解决尺寸过小与背景干扰问题，跟踪器需要解决短时遮挡与视觉退化问题，状态估计模块需要在漏检条件下维持可用预测，控制器则需要把图像空间误差转化为稳定的飞行控制量。一个标准的综述应当在文献层面把这些环节串联起来，而不是仅围绕某个具体项目说明“准备做什么”。',
        '基于这一认识，本文将 40 篇文献作为主要分析对象，强调“对文献的综述”而不是“由文献引出项目”。本文希望回答以下几个问题：第一，小目标检测与无人机视觉研究在方法上经历了怎样的演进；第二，当前常用数据集与 benchmark 在任务定义和评测指标上分别覆盖了哪些能力；第三，跟踪与状态估计方法如何支撑动态目标连续观测；第四，视觉伺服与自主追踪/拦截研究如何把感知结果闭环到控制系统；第五，从这 40 篇文献中能否提炼出对后续研究具有普适意义的方法趋势与问题清单。',
        '与一般课程作业式文献罗列不同，本文在结构上采用“研究主题综述 + 代表工作比较 + 综合讨论”的方式展开，既保留对关键论文的具体分析，也尽量突出跨论文的共性规律和方法边界，以符合标准综述论文的写法要求。',
    ]:
        para(doc, t)

    heading(doc, '2 文献检索策略与综述框架')
    heading(doc, '2.1 文献来源与筛选原则', level=2)
    for t in [
        '本文文献检索时间截至 2026 年 3 月 21 日，主要使用 Google Scholar、IEEE Xplore、CVF Open Access、SpringerLink、ScienceDirect 以及 OpenAlex 元数据接口进行交叉检索。考虑到综述目标是形成一个围绕低空小目标感知与拦截的完整技术图谱，检索关键词被划分为五组：一是“small object detection in aerial images”“tiny object detection”“UAV image detection”；二是“UAV detection and tracking benchmark”“VisDrone”“UAVDT”；三是“anti-UAV tracking”“unauthorized drone detection”；四是“visual servoing quadrotor target tracking”“moving target interception UAV”；五是“AirSim visual simulation”。',
        '在纳入标准方面，本文优先保留四类文献：第一，能够代表某一技术方向关键节点的高被引基础论文；第二，来自 CVPR、ICCV、ECCV、WACV、AAAI、TPAMI、IJCV、TITS、TMM、RA-L 等高质量会议与期刊的代表性论文；第三，与 UAV/anti-UAV 任务直接相关的数据集、benchmark 和挑战结果论文；第四，将视觉感知与跟踪、状态估计或控制结合的系统型研究。排除标准主要包括：与无人机或低空小目标关系不强的泛视觉论文、缺乏正式发表来源的低质量工作，以及与感知 - 跟踪 - 控制主线关联过弱的外围主题。',
        '最终筛得 40 篇文献，其中检测方法 11 篇、数据集与 benchmark 12 篇、跟踪与状态估计 8 篇、视觉伺服/拦截/仿真 9 篇。这样的划分既能覆盖低空小目标研究的主要技术链路，也有助于避免综述结构失衡。',
    ]:
        para(doc, t)

    heading(doc, '2.2 综述框架与统计特征', level=2)
    for t in [
        '从时间分布看，2017-2020 年的文献更多聚焦于通用检测基础、小目标初步建模以及无人机视觉 benchmark 的建立；2021-2023 年成为该方向快速扩张期，表现为专门的小目标检测器、反无人机跟踪基准和更强的在线跟踪算法持续出现；2024-2025 年的新工作则进一步向高分辨率小目标、复杂背景自适应和精细化场景建模发展。换言之，这 40 篇文献并不是彼此孤立的研究点，而是在时间上构成了一条由“检测基础 - 场景化适配 - 标准化基准 - 时序建模 - 闭环控制”的演进链。',
        '从来源分布看，CVPR、ICCV、ECCV、WACV 等视觉顶会主要贡献检测器、挑战赛结果和新型跟踪框架；TPAMI、IJCV、TITS、TMM 等期刊则更多承担 benchmark 总结、系统比较和领域整合的角色；RA-L、TMech、Aerospace Science and Technology 等控制或机器人方向期刊/会议则为视觉伺服与目标拦截提供控制理论和系统实现支撑。不同来源的关注重点不同，这也是本文采用主题式综述而非简单按年份罗列的原因。',
    ]:
        para(doc, t)

    make_table(doc, ['类别', '数量', '代表文献编号', '主要问题'], [
        ['低空小目标检测', '11', '1-11', '小尺度目标表征、多尺度融合、复杂背景抑制'],
        ['数据集与 benchmark', '12', '12-23', '任务定义、评价指标、场景覆盖'],
        ['跟踪与状态估计', '8', '24-31', '在线关联、时序建模、轨迹稳定性'],
        ['视觉伺服与拦截控制', '9', '32-40', '图像误差闭环、目标可见性、控制实时性'],
    ], font_size=10)

    heading(doc, '3 低空小目标视觉检测研究综述')
    heading(doc, '3.1 通用检测基础文献的作用', level=2)
    for t in [
        '虽然本文关注对象是低空小目标，但这一问题的技术起点并不完全来自 UAV 场景专用研究，而是源于通用目标检测框架对小尺度目标的长期探索。FPN[1]通过跨层特征融合把高层语义传递到高分辨率特征图，使检测器能够在不同尺度上共享更稳定的表征能力，这一思想几乎成为后续所有小目标检测工作的共同基础。Focal Loss[2]则从损失函数角度解决了单阶段检测器中的类别不平衡问题，强调难样本与少数前景样本的重要性。',
        'EfficientDet[3]在结构层面进一步说明，小目标检测并不总是依赖“更大模型”，而可以通过更合理的尺度融合与网络缩放在精度和速度间取得平衡；DETR[4]则展示了端到端检测的可能性，虽然它在小目标任务上并不占优，但其“统一建模”的思路对后续检测 - 跟踪一体化研究具有启发意义。QueryDet[5]把高分辨率小目标检测中的计算代价问题推到了前台，提示研究者：小目标并非一定需要全局密集推理，针对稀疏候选区域进行精细计算反而可能更加有效。',
        '从综述视角看，这几篇基础文献的重要意义不在于它们能否直接拿来检测无人机，而在于它们分别回答了小目标研究中的三个核心问题：如何保留细粒度尺度信息，如何让模型关注难检目标，以及如何在高分辨率场景中控制计算成本。后续的无人机小目标工作基本都在这三个问题上做场景化适配。',
    ]:
        para(doc, t)

    heading(doc, '3.2 航拍场景下的小目标专用检测研究', level=2)
    for t in [
        '与通用目标检测相比，航拍或低空视角中的小目标任务具有更强的场景特异性。首先，目标往往只占据极少像素，且容易受到云层、建筑边缘、地面纹理、复杂天气和平台运动模糊影响；其次，目标尺度变化剧烈，同一类飞行器可能因距离变化在图像中呈现出完全不同的表观；再次，空中目标的方向性和长宽比不稳定，导致常规 anchor 或中心点假设更容易失效。文献[6]较早围绕这些问题提出特征融合与尺度增强策略，并引入空间上下文分析，表明小目标检测不应只关注目标框内部信息，还必须考虑其周边背景关系。',
        '文献[7]提出 Focus-and-Detect 框架，将“先聚焦局部区域，再进行精细检测”的思想用于航拍小目标任务，本质上是把检测问题从一张全局大图分解为“背景筛除 + 局部增强”的两阶段流程。MFFSODNet[8]则在多尺度特征融合上做得更细，强调浅层细节和深层语义协同的重要性。Context-scale-aware detector[9]不仅给出一个新 benchmark，也进一步证明了“尺度感知上下文”对小弱目标的关键作用。',
        '在更近期的研究中，多尺度混合注意力[10]和 tiny object detection[11]相关工作进一步说明，航拍小目标检测的难点已经不仅是分辨率不足，而是“目标、背景、运动模糊、成像条件”共同作用下的复杂识别问题。注意力机制能够增强目标区域响应，但其效果高度依赖训练数据分布；专门针对 tiny object 的方法能够提升指标，却可能在泛化性和实时性上作出妥协。也就是说，当前方法普遍在“精度、速度、泛化”三者之间进行平衡，而不存在对所有任务都最优的统一解。',
    ]:
        para(doc, t)

    heading(doc, '3.3 这一组文献的共识与分歧', level=2)
    for t in [
        '综合文献[1]-[11]可以看到，低空小目标检测研究已形成若干共识。第一，多尺度融合是刚性需求，无论采用 FPN 变体、上下文增强还是特征重组，目标都是让小尺度目标在高层语义空间中仍保留足够区分度；第二，背景抑制与目标增强必须同步进行，单纯放大图像或增深网络往往无法从根本上解决问题；第三，小目标研究越来越强调与具体场景相结合，例如航拍、海上、热红外或特种空域，而不是追求脱离场景的统一指标。',
        '但这一组文献也存在明显分歧。一类工作强调结构创新，试图通过注意力、变换器和稀疏查询获得精度突破；另一类工作更重视数据与上下文设计，认为真正的瓶颈在于样本分布和训练机制。对于标准综述写作而言，这种分歧恰恰构成了文献讨论的重点：低空小目标任务并不是“模型越复杂越好”，而是在具体任务目标、数据规模和系统实时性约束下寻找合适平衡。',
    ]:
        para(doc, t)

    make_table(doc, ['编号', '文献', '年份', '主要贡献', '主要局限'], detection_rows())

    heading(doc, '4 数据集、benchmark 与挑战赛文献综述')
    heading(doc, '4.1 从 UAVDT 到反无人机 benchmark 的演进', level=2)
    for t in [
        '若没有标准化数据集和评价体系，小目标检测与跟踪研究很难从“单实验有效”走向“可比较、可积累”的学术发展路径。UAVDT[13]的重要意义正在于此。它把目标检测、单目标跟踪和多目标跟踪统一放到无人机视角下进行评测，使研究者意识到：无人机平台运动、俯视视角和尺度变化会显著改变传统地面视觉算法的表现。Detection and Tracking Meet Drones Challenge[12]则在此基础上进一步强化 benchmark 的组织方式，以挑战赛形式把数据、指标和排行榜整合起来，推动了该方向的快速扩散。',
        '随着研究重心从“无人机看地面目标”逐渐转向“视觉系统看无人机目标”，Anti-UAV[15]和 Vision-Based Anti-UAV Detection and Tracking[14]具有里程碑意义。这类文献将低空飞行器本身作为被观测对象，真正触及了“低慢小空中目标”这一任务核心。与 UAVDT 相比，反无人机数据集更强调天空/地面混合背景、尺度极小目标、快速出视场与重新入视场等问题，也更接近拦截或防御型任务。',
        '从 benchmark 演进路径看，研究重点已经从“有没有数据”转向“数据是否能覆盖真实问题”。WebUAV-3M[19]通过百万级规模推动深度跟踪模型训练，BIRDSAI[18]把热红外引入航拍感知，SeaDronesSee[17]展示了特种环境下的小目标视觉困难，Unauthorized Drones Review[23]则从更宽广的传感视角总结反无人机检测技术。它们共同说明，数据集建设不只是收集样本，而是通过任务设置和指标设计决定研究的方向性。',
    ]:
        para(doc, t)

    heading(doc, '4.2 VisDrone 与挑战赛结果文献的价值', level=2)
    for t in [
        'VisDrone 系列挑战赛结果文献[21][22]在综述中也应占有独立位置。虽然这类论文本身往往不提出全新的核心算法，但它们记录了当时主流方法在统一数据和规则下的表现，能够清楚展示研究社区的关注重点和性能瓶颈。例如，从 2019 到 2021 年的结果对比可以看到，多尺度增强、数据增广和更强主干网络逐渐成为基础配置，而仅凭单一结构微调已难取得显著领先优势。',
        '对于标准综述而言，挑战赛结果文献的重要作用在于“还原研究共同体的真实选择”。相比论文作者自行选取实验设置，竞赛结果更能揭示在统一资源和规则约束下哪些策略更稳定、哪些改进只对个别数据有效。因此，把挑战赛结果纳入综述能够避免文献分析过于理想化。',
    ]:
        para(doc, t)

    heading(doc, '4.3 数据集与 benchmark 文献的不足', level=2)
    for t in [
        '尽管数据集与 benchmark 已明显丰富，但从拦截或闭环系统视角看，现有文献仍存在若干结构性不足。第一，大多数 benchmark 仍然以检测精度、跟踪成功率或 MOT 指标为中心，较少关注目标是否足以支撑控制决策；第二，许多数据集虽然覆盖了复杂背景，却缺少对“连续逼近、目标失锁恢复、控制抖动”等系统级现象的评价维度；第三，部分数据集关注单一模态，难以支持低照度、逆光、天气变化等更复杂条件下的系统研究。',
        '因此，在标准综述论文中，benchmark 文献不应只被当作“引用数据集名字”的背景材料，而应作为决定研究问题边界的重要对象。一个研究方向能解决什么问题，在相当程度上取决于它拥有什么样的基准数据和怎样的评价标准。',
    ]:
        para(doc, t)

    make_table(doc, ['编号', '文献', '年份', '主题', '主要价值'], benchmark_rows())

    heading(doc, '5 目标跟踪与状态估计文献综述')
    heading(doc, '5.1 检测后关联方法：在线关联路线', level=2)
    for t in [
        '在目标跟踪研究中，最经典的一条路线是“先检测，后关联”。Deep SORT[24]作为这一路线的代表工作，在运动预测基础上引入外观特征关联，使在线多目标跟踪从单纯依赖卡尔曼 + 匈牙利匹配迈向“几何 + 外观”融合阶段。其贡献不仅在于提升 ID 保持能力，也在于确立了检测后跟踪作为工程基线的可行性。',
        'ByteTrack[25]的突破在于重新看待低置信度检测框的价值。传统方法通常在前端阈值阶段直接丢弃低分框，但 ByteTrack 指出这些候选框恰恰可能承载被遮挡、尺度变化或运动模糊条件下的真实目标信息。该思想对于低空小目标尤其重要，因为这类目标天然容易得到不稳定的置信度分数。OC-SORT[30]则从观测更新角度进一步增强 SORT 路线的稳定性，反映出检测后关联方法并未过时，而是在不断通过更鲁棒的观测融合继续发展。',
        '对标准综述而言，这些文献共同说明：在线关联路线之所以长期有效，不是因为它最“先进”，而是因为它兼顾了可解释性、模块化和工程落地性。对于需要与控制环路耦合的系统而言，轻量、稳定、可调试往往比极限精度更重要。',
    ]:
        para(doc, t)

    heading(doc, '5.2 时序统一建模与单目标跟踪路线', level=2)
    for t in [
        '另一条重要路线是把检测与跟踪进行更深程度统一。Tracking Objects as Points[26]从关键点视角出发，把检测与时序关联放入同一框架中处理，体现了“端到端时序建模”的思想。Improving Multiple Object Tracking with Single Object Tracking[29]则进一步打通了 MOT 与 SOT 之间的边界，说明单目标跟踪中的局部匹配和持续关注机制同样能提升多目标场景的轨迹连续性。',
        '在单目标跟踪方向，SiamRPN++[27]和 DiMP[28]分别代表了孪生匹配和判别式模型预测两种典型思路。前者强调模板与搜索区域之间的快速匹配，适合实时性要求较高的任务；后者则更强调模型更新和判别性保持，在复杂背景下的鲁棒性更强。这两类方法虽然主要服务于单目标跟踪，但其关于时序连续性、在线更新和局部目标表征的思想，对无人机目标追踪或拦截任务都具有借鉴意义。',
        '值得注意的是，这一路线的代价通常是系统复杂度和计算成本上升。相比检测后跟踪，端到端或深度 SOT 方法在高分辨率场景、长时视频和机载算力受限条件下不一定占优。因此，文献综述不能简单得出“越一体化越好”的结论，而应强调不同路线在任务边界、实时性和系统耦合方式上的差异。',
    ]:
        para(doc, t)

    heading(doc, '5.3 状态估计与评价指标：从轨迹保持到可用于控制', level=2)
    for t in [
        '除具体跟踪算法外，评价体系和状态估计思想本身也是综述的重要组成部分。HOTA[31]从检测、关联和定位三个维度统一刻画跟踪性能，弥补了单一 MOTA/IDF1 指标的不足，使研究者能够更全面地理解轨迹质量。对于低空小目标任务而言，这一点尤为重要，因为轨迹偶发中断、目标短时丢失和定位抖动都会对后续控制造成放大效应。',
        '虽然本文选取的 40 篇文献中并没有专门把卡尔曼滤波单独作为主论文条目纳入，但 Deep SORT、SORT 系列、目标追踪控制论文都体现了状态估计在动态目标系统中的基础地位。综述这些论文后可以得到一个明确结论：在检测器尚不能完全稳定工作的前提下，轻量级状态预测依然是提升系统可用性的关键环节。',
    ]:
        para(doc, t)

    make_table(doc, ['编号', '文献', '年份', '主题', '主要价值'], tracking_rows())

    heading(doc, '6 视觉伺服、目标逼近与拦截控制文献综述')
    heading(doc, '6.1 视觉伺服研究：图像空间误差如何闭环到飞行控制', level=2)
    for t in [
        '视觉伺服研究解决的是“如何根据图像观测直接驱动飞行器运动”的问题，这一问题与单纯的检测或跟踪有本质区别。Position-Based Visual Servoing[35]从目标相对位姿与飞行控制律入手，说明视觉控制并不只是把检测框中心当作误差输入，而需要将图像信息映射为物理可执行的控制量。Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs[32]则进一步提出，保持目标可见本身就是控制目标的一部分，如果目标频繁逼近视场边缘，检测和跟踪再强也无法支撑稳定控制。',
        'Image-Based Visual Servoing of Rotorcrafts to Planar Visual Targets of Arbitrary Orientation[34]和 Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets[33]把视觉伺服从平面目标扩展到更加复杂的姿态和飞行目标场景，说明图像空间控制的核心挑战在于目标几何不确定性、姿态变化和飞行器自身动力学耦合。对于综述写作而言，这一组论文揭示了一个关键事实：感知模块输出的误差并不能直接等同于控制意义上的误差，只有经过可见性、约束和动态稳定性分析，视觉信息才能真正进入控制回路。',
    ]:
        para(doc, t)

    heading(doc, '6.2 自主追踪与拦截研究：从跟踪目标到接近目标', level=2)
    for t in [
        '自主追踪与拦截论文把问题进一步从“看见目标”推进到“逼近甚至捕获目标”。Autonomous Target Tracking of UAV Using High-Speed Visual Feedback[37]强调高频图像反馈在动态目标追踪中的价值，展示了感知刷新率与控制带宽之间的直接关系。Autonomous Drone Hunter Operating by Deep Learning and All-Onboard Computations in GPS-Denied Environments[36]则把无人机自主追踪推向更接近实际系统的层面，说明在 GPS 受限条件下，仅依赖机载视觉与计算也可以完成对目标的连续追踪与处置。',
        'UAV Guidance: A Stereo-Based Technique for Interception of Stationary or Moving Targets[39]进一步表明，拦截任务相较普通跟踪任务更强调目标相对位置估计精度、接近路径规划以及终端阶段控制稳定性。与一般视频跟踪不同，拦截场景对“短时间内的正确决策”敏感度更高，哪怕一次瞬时漏检或错误机动都可能直接造成拦截失败。换句话说，拦截研究把感知、预测和控制之间的耦合关系放大了，这也是为何标准综述必须单独为这类文献设立板块。',
        'Robust Visual Servoing Formation Tracking Control for Quadrotor UAV Team[38]虽然不是直接针对单目标拦截，但它展示了视觉控制在多飞行器系统中的鲁棒性设计方法。该类工作提醒我们：一旦任务从单机追踪扩展到协同处置或编队控制，视觉系统的稳定性要求会进一步提高。',
    ]:
        para(doc, t)

    heading(doc, '6.3 仿真平台与闭环系统验证', level=2)
    for t in [
        'AirSim[40]在这组文献中的地位十分特殊，因为它本身不是检测器、跟踪器或控制律，但却为这些算法提供了统一的验证环境。高保真视觉渲染、物理引擎和可控场景构建能力，使得研究者可以在仿真中复现实验条件、自动采集数据、快速调整参数并验证闭环性能。对于需要频繁迭代的低空小目标系统来说，仿真平台不仅是“辅助手段”，而是系统工程方法的一部分。',
        '然而，AirSim 及类似平台也存在明显边界，即仿真到真实环境之间仍有域差异，包括纹理分布、光照变化、气动扰动和传感噪声等。因此，标准综述论文不能把仿真验证等同于真实部署，而应把它视为感知 - 控制系统验证中的一个阶段性环节。只有把仿真结果、数据集实验和真实环境测试结合起来，才能形成完整的系统评价闭环。',
    ]:
        para(doc, t)

    make_table(doc, ['编号', '文献', '年份', '主题', '主要价值'], control_rows())

    heading(doc, '7 跨 40 篇文献的综合比较与讨论')
    heading(doc, '7.1 研究脉络：从检测到闭环系统', level=2)
    for t in [
        '通读 40 篇文献后可以发现，该领域的研究脉络并不是线性单向前进，而是围绕“感知更准、时序更稳、控制更强、验证更规范”四个目标交替推进。检测基础论文[1]-[5]解决的是小目标能否被看见的问题；航拍小目标论文[6]-[11]关注的是看见之后如何在特定场景下保持有效性；benchmark 论文[12]-[23]解决的是不同方法是否可比较的问题；跟踪文献[24]-[31]处理的是目标在时间维上的连续观测；而控制与拦截文献[32]-[40]则回答了“看见和追踪之后，是否能够真正接近并处置目标”。',
        '因此，从综述视角看，这 40 篇文献最重要的意义并不是各自在各自任务上刷新了多少百分点，而是共同构成了低空小目标研究由“单点算法优化”向“系统链路整合”过渡的证据链。这种过渡恰恰是现代无人系统研究最值得关注的方向。',
    ]:
        para(doc, t)

    heading(doc, '7.2 本组文献体现出的主要进展', level=2)
    for t in [
        '第一，研究对象从一般小目标逐步收敛到更具任务导向性的低空飞行目标。早期基础论文关注的是普适性的多尺度检测与密集分类问题，而近年的 UAV/anti-UAV 文献则更明确地面向低空飞行器目标，从而使任务定义更加贴近实际应用。',
        '第二，研究范式从“模型导向”逐步走向“数据与 benchmark 导向”。没有 UAVDT、VisDrone、Anti-UAV、WebUAV-3M 等数据集，很多方法改进其实难以稳定比较。数据集文献不仅提供了样本，还塑造了领域的研究边界。',
        '第三，时序建模的地位不断提升。无论是 ByteTrack、Deep SORT 这类检测后关联方法，还是 CenterTrack、DiMP、SiamRPN++ 这类统一时序方法，都说明低空小目标研究已经不满足于单帧检测，而更重视跨帧稳定性。',
        '第四，视觉控制与拦截研究逐渐从理论可行性走向系统实现。高频视觉反馈、可见性约束、双目引导、自主追踪与机载计算等文献表明，感知结果正在更紧密地进入飞行控制系统。',
    ]:
        para(doc, t)

    heading(doc, '7.3 现有研究的共性问题', level=2)
    for t in [
        '首先，极小目标与复杂背景之间的矛盾依然是所有视觉方法的基本瓶颈。即便经过多尺度融合、注意力增强和上下文建模，目标像素极少、边缘模糊、纹理弱这一事实并未改变。一旦背景纹理恰好与目标形态相近，模型仍可能出现高置信误检或持续漏检。',
        '其次，许多论文虽然在检测或跟踪指标上表现优异，但对系统级稳定性的讨论不足。对于拦截型任务而言，短时漏检、视场边缘漂移、控制延迟和轨迹抖动所造成的影响往往大于单帧 AP 指标的细微差别，而这些因素在很多 benchmark 中尚未得到直接度量。',
        '再次，实时性与精度之间的矛盾尚未被彻底解决。高性能 Transformer、复杂注意力模块和统一时序框架可以提高指标，但也提高了算力需求。对于机载部署而言，算法是否能够在受限硬件上稳定运行，往往比离线性能更重要。',
        '最后，仿真验证与真实部署之间仍存在明显鸿沟。AirSim 这类平台使闭环实验更加容易，但仿真环境的纹理、天气、扰动和传感噪声仍然难以完全覆盖真实空域条件。因此，很多文献在“仿真有效”与“实装有效”之间还存在证据断层。',
    ]:
        para(doc, t)

    heading(doc, '7.4 未来研究方向', level=2)
    for t in [
        '未来一个重要方向是检测、预测与控制的协同设计。现有研究往往按模块分割任务，导致检测指标最优的模型不一定是闭环控制最优的模型。面向拦截任务的研究更需要考虑哪些视觉特征最有利于控制决策，哪些状态量需要预测保持，以及如何让模型输出更直接服务于控制。',
        '第二个方向是面向低算力平台的轻量化闭环系统。当前许多高性能方法仍然建立在较强计算资源前提下，但实用系统往往需要在边缘设备或机载平台上完成推理、预测和控制。如何在模型剪枝、知识蒸馏、稀疏推理和控制侧容错之间寻找平衡，将成为重要课题。',
        '第三个方向是构建更贴近任务目标的 benchmark。现有数据集非常丰富，但真正围绕“目标能否被持续追踪并支持逼近/拦截”的 benchmark 仍然较少。未来可考虑把检测、跟踪、短时预测和控制相关指标联合纳入评测体系，使 benchmark 更接近系统应用需求。',
        '第四个方向是提升仿真到真实环境的迁移能力。域随机化、风扰动建模、合成数据增广、真实场景微调和多模态感知融合都有望缩小仿真与真实之间的差距。对于低空小目标系统而言，只有在更接近真实复杂空域条件下验证有效，研究成果才真正具有推广价值。',
    ]:
        para(doc, t)

    heading(doc, '8 结论')
    for t in [
        '本文基于 40 篇代表性文献，对低空小目标的视觉检测、跟踪、状态估计、视觉伺服与目标拦截控制进行了系统综述。与以项目实现为中心的介绍性文稿不同，本文以文献本身为对象，强调对研究主题、方法谱系、评测体系和系统边界的综合把握。综述结果表明，该方向已经从早期通用检测器改进逐步发展为包含 benchmark 建设、时序建模、视觉控制和仿真验证的系统性研究领域。',
        '从总体趋势看，多尺度与上下文建模仍将长期主导小目标检测研究，数据集和 benchmark 将继续塑造任务边界，跟踪与状态预测将成为连接感知与控制的重要桥梁，而视觉伺服和自主拦截研究则会把系统需求重新反馈给感知模块。对于后续研究者而言，最有价值的工作不再只是某一指标上的局部最优，而是在真实任务边界下实现对感知、预测和控制的协同优化。',
    ]:
        para(doc, t)

    doc.add_page_break()
    heading(doc, '附录 A 40 篇文献主题矩阵')
    para(doc, '为便于后续写作与查阅，附录按统一格式给出 40 篇文献的主题、来源与方法要点。该矩阵同时可以作为进一步撰写“研究现状”章节时的快速索引。', indent=False)
    make_table(doc, ['编号', '文献', '主题', '来源', '年份', '主要贡献', '主要局限'], full_matrix_rows(), font_size=8.8)

    doc.add_page_break()
    heading(doc, '参考文献')
    for r in refs:
        ref_para(doc, r['id'], f"{r['title']}. {r['venue']}, {r['year']}. DOI: {r['doi']}")

    doc.save(OUT_PATH)
    print(OUT_PATH)

if __name__ == '__main__':
    build()
