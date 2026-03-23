from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r'D:\Xstarlab\UEProjects\GraduationProject\GraduationProject')
OUT_DIR = ROOT / 'output' / 'doc'
OUT_DIR.mkdir(parents=True, exist_ok=True)

TASK_PATH = OUT_DIR / '毕业设计任务书与论文计划.docx'
REVIEW_PATH = OUT_DIR / '文献综述-基于机器视觉的低慢小目标拦截技术研究.docx'

TITLE = '基于机器视觉的低慢小目标拦截技术研究'
STUDENT = '颜飞翔'
STUDENT_ID = '2229409056'
DATE_STR = '2026年3月21日'

references = [
    {'id': 1, 'title': 'Feature Pyramid Networks for Object Detection', 'venue': 'CVPR', 'year': '2017', 'doi': '10.1109/CVPR.2017.106'},
    {'id': 2, 'title': 'Focal Loss for Dense Object Detection', 'venue': 'ICCV', 'year': '2017', 'doi': '10.1109/ICCV.2017.324'},
    {'id': 3, 'title': 'EfficientDet: Scalable and Efficient Object Detection', 'venue': 'CVPR', 'year': '2020', 'doi': '10.1109/CVPR42600.2020.01079'},
    {'id': 4, 'title': 'End-to-End Object Detection with Transformers', 'venue': 'ECCV', 'year': '2020', 'doi': '10.1007/978-3-030-58452-8_13'},
    {'id': 5, 'title': 'QueryDet: Cascaded Sparse Query for Accelerating High-Resolution Small Object Detection', 'venue': 'CVPR', 'year': '2022', 'doi': '10.1109/CVPR52688.2022.01330'},
    {'id': 6, 'title': 'Small Object Detection in Unmanned Aerial Vehicle Images Using Feature Fusion and Scaling-Based Single Shot Detector With Spatial Context Analysis', 'venue': 'IEEE TCSVT', 'year': '2019', 'doi': '10.1109/TCSVT.2019.2905881'},
    {'id': 7, 'title': 'Focus-and-Detect: A Small Object Detection Framework for Aerial Images', 'venue': 'Signal Processing: Image Communication', 'year': '2022', 'doi': '10.1016/j.image.2022.116675'},
    {'id': 8, 'title': 'MFFSODNet: Multiscale Feature Fusion Small Object Detection Network for UAV Aerial Images', 'venue': 'IEEE TIM', 'year': '2024', 'doi': '10.1109/TIM.2024.3381272'},
    {'id': 9, 'title': 'A Context-Scale-Aware Detector and a New Benchmark for Remote Sensing Small Weak Object Detection in Unmanned Aerial Vehicle Images', 'venue': 'JAG', 'year': '2022', 'doi': '10.1016/j.jag.2022.102966'},
    {'id': 10, 'title': 'Small Object Detection in Unmanned Aerial Vehicle Images Using Multi-Scale Hybrid Attention', 'venue': 'Engineering Applications of Artificial Intelligence', 'year': '2023', 'doi': '10.1016/j.engappai.2023.107455'},
    {'id': 11, 'title': 'Tiny Object Detection in Aerial Images', 'venue': 'ICPR', 'year': '2021', 'doi': '10.1109/ICPR48806.2021.9413340'},
    {'id': 12, 'title': 'Detection and Tracking Meet Drones Challenge', 'venue': 'IEEE TPAMI', 'year': '2021', 'doi': '10.1109/TPAMI.2021.3119563'},
    {'id': 13, 'title': 'The Unmanned Aerial Vehicle Benchmark: Object Detection, Tracking and Baseline', 'venue': 'IJCV', 'year': '2019', 'doi': '10.1007/s11263-019-01266-1'},
    {'id': 14, 'title': 'Vision-Based Anti-UAV Detection and Tracking', 'venue': 'IEEE TITS', 'year': '2022', 'doi': '10.1109/TITS.2022.3177627'},
    {'id': 15, 'title': 'Anti-UAV: A Large-Scale Benchmark for Vision-Based UAV Tracking', 'venue': 'IEEE TMM', 'year': '2021', 'doi': '10.1109/TMM.2021.3128047'},
    {'id': 16, 'title': 'A Unified Transformer-Based Tracker for Anti-UAV Tracking', 'venue': 'CVPRW', 'year': '2023', 'doi': '10.1109/CVPRW59228.2023.00305'},
    {'id': 17, 'title': 'SeaDronesSee: A Maritime Benchmark for Detecting Humans in Open Water', 'venue': 'WACV', 'year': '2022', 'doi': '10.1109/WACV51458.2022.00374'},
    {'id': 18, 'title': 'BIRDSAI: A Dataset for Detection and Tracking in Aerial Thermal Infrared Videos', 'venue': 'WACV', 'year': '2020', 'doi': '10.1109/WACV45572.2020.9093284'},
    {'id': 19, 'title': 'WebUAV-3M: A Benchmark for Unveiling the Power of Million-Scale Deep UAV Tracking', 'venue': 'IEEE TPAMI', 'year': '2022', 'doi': '10.1109/TPAMI.2022.3232854'},
    {'id': 20, 'title': 'Unmanned Aerial Vehicle Visual Detection and Tracking Using Deep Neural Networks: A Performance Benchmark', 'venue': 'ICCVW', 'year': '2021', 'doi': '10.1109/ICCVW54120.2021.00142'},
    {'id': 21, 'title': 'VisDrone-DET2019: The Vision Meets Drone Object Detection in Image Challenge Results', 'venue': 'ICCVW', 'year': '2019', 'doi': '10.1109/ICCVW.2019.00030'},
    {'id': 22, 'title': 'VisDrone-DET2021: The Vision Meets Drone Object Detection Challenge Results', 'venue': 'ICCVW', 'year': '2021', 'doi': '10.1109/ICCVW54120.2021.00319'},
    {'id': 23, 'title': 'On the Detection of Unauthorized Drones - Techniques and Future Perspectives: A Review', 'venue': 'IEEE Sensors Journal', 'year': '2022', 'doi': '10.1109/JSEN.2022.3171293'},
    {'id': 24, 'title': 'Simple Online and Realtime Tracking with a Deep Association Metric', 'venue': 'ICIP', 'year': '2017', 'doi': '10.1109/ICIP.2017.8296962'},
    {'id': 25, 'title': 'ByteTrack: Multi-Object Tracking by Associating Every Detection Box', 'venue': 'ECCV', 'year': '2022', 'doi': '10.1007/978-3-031-20047-2_1'},
    {'id': 26, 'title': 'Tracking Objects as Points', 'venue': 'ECCV', 'year': '2020', 'doi': '10.1007/978-3-030-58548-8_28'},
    {'id': 27, 'title': 'SiamRPN++: Evolution of Siamese Visual Tracking With Very Deep Networks', 'venue': 'CVPR', 'year': '2019', 'doi': '10.1109/CVPR.2019.00441'},
    {'id': 28, 'title': 'Learning Discriminative Model Prediction for Tracking', 'venue': 'ICCV', 'year': '2019', 'doi': '10.1109/ICCV.2019.00628'},
    {'id': 29, 'title': 'Improving Multiple Object Tracking with Single Object Tracking', 'venue': 'CVPR', 'year': '2021', 'doi': '10.1109/CVPR46437.2021.00248'},
    {'id': 30, 'title': 'Observation-Centric SORT: Rethinking SORT for Robust Multi-Object Tracking', 'venue': 'CVPR', 'year': '2023', 'doi': '10.1109/CVPR52729.2023.00934'},
    {'id': 31, 'title': 'HOTA: A Higher Order Metric for Evaluating Multi-Object Tracking', 'venue': 'IJCV', 'year': '2020', 'doi': '10.1007/s11263-020-01375-2'},
    {'id': 32, 'title': 'Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs', 'venue': 'IEEE/ASME TMech', 'year': '2019', 'doi': '10.1109/TMECH.2019.2906430'},
    {'id': 33, 'title': 'Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets', 'venue': 'IEEE RA-L', 'year': '2023', 'doi': '10.1109/LRA.2023.3245416'},
    {'id': 34, 'title': 'Image-Based Visual Servoing of Rotorcrafts to Planar Visual Targets of Arbitrary Orientation', 'venue': 'IEEE RA-L', 'year': '2021', 'doi': '10.1109/LRA.2021.3101878'},
    {'id': 35, 'title': 'Position-Based Visual Servoing for Target Tracking by a Quadrotor UAV', 'venue': 'AIAA GNC', 'year': '2016', 'doi': '10.2514/6.2016-2092'},
    {'id': 36, 'title': 'Autonomous Drone Hunter Operating by Deep Learning and All-Onboard Computations in GPS-Denied Environments', 'venue': 'PLOS ONE', 'year': '2019', 'doi': '10.1371/journal.pone.0225092'},
    {'id': 37, 'title': 'Autonomous Target Tracking of UAV Using High-Speed Visual Feedback', 'venue': 'Applied Sciences', 'year': '2019', 'doi': '10.3390/app9214552'},
    {'id': 38, 'title': 'Robust Visual Servoing Formation Tracking Control for Quadrotor UAV Team', 'venue': 'Aerospace Science and Technology', 'year': '2020', 'doi': '10.1016/j.ast.2020.106061'},
    {'id': 39, 'title': 'UAV Guidance: A Stereo-Based Technique for Interception of Stationary or Moving Targets', 'venue': 'Lecture Notes in Computer Science', 'year': '2015', 'doi': '10.1007/978-3-319-22416-9_30'},
    {'id': 40, 'title': 'AirSim: High-Fidelity Visual and Physical Simulation for Autonomous Vehicles', 'venue': 'Springer FSR', 'year': '2017', 'doi': '10.1007/978-3-319-67361-5_40'},
]


def set_east_asia(run, font_name='宋体'):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(12)


def add_title(doc: Document, text: str, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    set_east_asia(r, '黑体')
    return p


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    set_east_asia(r, '宋体')
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12)
    set_east_asia(r, '黑体')
    return p


def add_body(doc: Document, text: str, first_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(12)
    set_east_asia(r, '宋体')
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(12)
    set_east_asia(r, '宋体')
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(12)
    set_east_asia(r, '宋体')
    return p


def add_reference(doc: Document, idx: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    r = p.add_run(f'[{idx}] {text}')
    r.font.size = Pt(10.5)
    set_east_asia(r, '宋体')
    return p


def add_table_info(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        left = table.cell(i, 0)
        right = table.cell(i, 1)
        left.text = k
        right.text = v
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell in (left, right):
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.25
                for run in para.runs:
                    run.font.size = Pt(11)
                    set_east_asia(run, '宋体')
    return table


def build_task_doc():
    doc = Document()
    configure_document(doc)
    add_title(doc, '毕业设计任务书与论文计划', 18)
    add_subtitle(doc, TITLE)
    doc.add_paragraph()
    add_table_info(doc, [
        ('课题名称', TITLE),
        ('学生姓名', STUDENT),
        ('学号', STUDENT_ID),
        ('文档日期', DATE_STR),
        ('项目基础', 'UE/AirSim 仿真 + YOLO 检测 + Kalman 预测 + 视觉引导拦截控制'),
    ])

    add_heading(doc, '一、设计（论文）的主要任务及目标')
    for para in [
        '本课题围绕低慢小目标在复杂低空环境中的视觉检测、跟踪预测与拦截控制问题展开，依托现有 Unreal Engine/AirSim 仿真平台与 Python 感知控制链路，构建一套从场景构建、图像采集、数据标注、检测训练到控制验证的完整研究流程。课题拟解决小目标在复杂背景下易漏检、跟踪不连续、控制耦合弱等问题，为低成本、可复现实验条件下的无人机目标拦截研究提供技术支撑。',
        '论文目标不是单点算法优化，而是形成一条可以闭环运行的系统链路：利用仿真环境生成目标与拦截器运动场景，构建低慢小目标数据集；采用 YOLO 类模型完成目标检测；基于图像特征构建卡尔曼预测模块，实现目标短时状态估计；设计以图像误差和面积反馈为核心的视觉引导控制策略，并通过多组仿真实验评价拦截效果。',
    ]:
        add_body(doc, para)
    for item in [
        '搭建面向低慢小目标拦截的仿真测试平台，实现目标与拦截器的场景生成、运动控制和图像采集。',
        '完成低慢小目标图像数据采集与自动标注，构建可用于检测训练的数据集。',
        '研究适用于低慢小目标的视觉检测方法，训练并优化 YOLO 检测模型，提高复杂背景下的小目标识别能力。',
        '研究基于视觉特征的目标状态估计方法，引入卡尔曼预测提升目标短时运动预测与丢失恢复能力。',
        '设计基于图像误差和面积反馈的视觉引导拦截控制方法，实现目标搜索、跟踪、逼近与拦截过程控制。',
        '通过多组仿真实验对系统在检测精度、跟踪稳定性、响应速度和拦截成功率等方面进行测试与分析。',
        '完成毕业论文、文献综述和答辩材料撰写。',
    ]:
        add_number(doc, item)

    add_heading(doc, '二、设计（论文）的主要内容')
    for item in [
        '低慢小目标拦截仿真环境构建：基于 Unreal Engine 与 AirSim 建立无人机拦截实验平台，完成拦截器、目标机和引导对象配置。',
        '数据采集与数据集构建：设计目标机与拦截器相对运动模式，利用仿真平台采集图像并自动生成标注。',
        '低慢小目标视觉检测研究：围绕小目标尺寸小、纹理弱、背景复杂等特点，完成检测模型训练与效果分析。',
        '目标跟踪与运动状态预测研究：根据检测框中心、尺度和时间变化引入卡尔曼预测，提升漏检情况下的控制稳定性。',
        '视觉引导拦截控制研究：设计搜索、跟踪、逼近和捕获状态机，利用 PID 和视觉误差完成闭环控制。',
        '系统集成与实验分析：设置不同目标运动模式、不同检测条件和不同控制参数，对系统性能进行对比验证。',
    ]:
        add_number(doc, item)

    add_heading(doc, '三、设计（论文）的基本要求')
    for item in [
        '系统应能够在仿真环境下完成低慢小目标的搜索、检测、跟踪和拦截过程验证。',
        '完成数据采集、标注、训练与测试流程，形成可复现的实验链路。',
        '检测模型应输出目标位置、置信度和尺度信息，并具备一定实时性。',
        '状态预测模块应能在短时漏检情况下维持控制连续性。',
        '建立实验评价指标，对检测效果、跟踪稳定性和拦截成功率进行分析。',
        '论文写作应符合本科毕业设计规范，结构完整、图表规范、引用准确。',
    ]:
        add_number(doc, item)

    add_heading(doc, '四、文献检索与综述写作计划')
    add_body(doc, '文献检索围绕“低慢小目标检测、无人机目标跟踪、反无人机基准、视觉伺服与引导拦截、仿真验证”五条技术线展开。检索数据库包括 Google Scholar、IEEE Xplore、ACM Digital Library、CVF Open Access、SpringerLink、ScienceDirect 和 arXiv。纳入文献以 SCI、EI、CVPR、ICCV、ECCV、WACV、TPAMI、IJCV、TITS、TMM 等高质量来源为主。')
    for item in [
        '第 1 阶段：检索并筛选 40 篇左右高质量文献，建立分主题阅读清单。',
        '第 2 阶段：按“小目标检测、数据集与基准、跟踪预测、视觉伺服与拦截、仿真平台”完成主题归纳。',
        '第 3 阶段：形成文献综述初稿，突出研究脉络、代表方法、现有不足及本文课题切入点。',
        '第 4 阶段：结合项目实现，补充系统结构图、算法流程图和实验指标说明。',
    ]:
        add_number(doc, item)

    add_heading(doc, '五、进度安排')
    schedule = [
        ('第 1-2 周', '明确课题需求，阅读相关文献，完成任务书与开题准备。'),
        ('第 3-5 周', '完成仿真环境梳理、场景配置和数据采集流程设计。'),
        ('第 6-8 周', '完成低慢小目标数据采集、数据清洗与数据集构建。'),
        ('第 9-11 周', '完成目标检测模型训练、调参与检测实验分析。'),
        ('第 12-13 周', '完成目标状态预测与视觉引导拦截控制算法设计与调试。'),
        ('第 14-15 周', '完成系统联调与拦截实验，整理实验结果与图表。'),
        ('第 16 周', '完成文献综述与论文初稿撰写。'),
        ('第 17 周', '完成论文修改、定稿及答辩准备。'),
    ]
    for week, content in schedule:
        add_body(doc, f'{week}：{content}', first_indent=False)

    add_heading(doc, '六、任务书主要参考文献')
    for ref in references[:15]:
        add_reference(doc, ref['id'], f"{ref['title']}. {ref['venue']}, {ref['year']}. DOI: {ref['doi']}")

    doc.save(TASK_PATH)


def build_review_doc():
    doc = Document()
    configure_document(doc)
    add_title(doc, TITLE, 18)
    add_subtitle(doc, '文献综述')
    add_subtitle(doc, f'作者：{STUDENT}    学号：{STUDENT_ID}')

    add_heading(doc, '摘要')
    add_body(doc, '低慢小目标通常具有飞行高度低、运动速度慢、目标尺寸小和背景干扰强等特征，典型目标包括微小型无人机、低空飞行器和其他难以稳定感知的空中目标。围绕该类目标的检测、跟踪、状态估计与拦截控制，本文在前期筛选的 40 篇高质量文献基础上，从小目标检测方法、无人机与反无人机数据集及基准、目标跟踪与运动状态估计、视觉伺服与引导拦截控制、仿真验证平台五个方面进行了系统综述。综述结果表明：低慢小目标研究已由单一检测模型改进逐步转向“检测 - 跟踪 - 预测 - 控制”一体化方向；多尺度特征建模、长时关联能力、视觉可见性约束和仿真闭环验证已成为影响系统性能的关键因素。结合现有毕设项目基础，本文进一步归纳了面向仿真闭环拦截系统的研究切入点，为后续论文撰写与系统完善提供理论支撑。')

    add_heading(doc, '关键词')
    add_body(doc, '低慢小目标；机器视觉；目标检测；视觉跟踪；卡尔曼预测；视觉伺服；无人机拦截', first_indent=False)

    add_heading(doc, '1 引言')
    for para in [
        '低慢小目标因飞行高度低、速度慢、尺寸小和回波弱，往往对传统预警与拦截系统构成挑战。在视觉观测条件下，这类目标在图像中通常只占据少量像素，容易受到光照变化、背景杂波、遮挡与运动模糊影响，从而导致检测误差和状态估计不稳定。随着低空无人系统应用范围不断扩大，如何依托机器视觉手段构建对低慢小目标的感知、预测与拦截闭环，逐步成为智能无人系统研究中的重要方向。',
        '从技术链路上看，低慢小目标拦截并非单一检测问题，而是一个涉及场景构建、数据采集、目标检测、目标跟踪、运动预测、控制决策和实验评估的系统工程。现有研究中，一部分工作聚焦于小目标检测器结构改进，以增强微小目标的可分辨性；另一部分工作围绕无人机视觉数据集、单目标或多目标跟踪、反无人机基准展开，推动了算法评测标准化；还有一部分研究将视觉伺服与目标逼近控制结合，用于动态目标追踪和拦截。',
        '对于本文毕设对应的项目实现而言，系统已经具备 AirSim 仿真环境、图像采集与自动标注、YOLO 检测训练、Kalman 预测以及视觉引导拦截控制的基础模块，因此有必要对相关文献进行系统梳理，明确各类研究成果的技术定位、适用场景与不足，并在此基础上形成与本课题实现相匹配的综述框架。',
    ]:
        add_body(doc, para)

    add_heading(doc, '2 文献筛选方法与分类原则')
    for para in [
        '本文文献检索时间截至 2026 年 3 月 21 日，检索来源主要包括 Google Scholar、IEEE Xplore、CVF Open Access、SpringerLink、ScienceDirect 和 OpenAlex 元数据接口。检索关键词包括“small object detection in aerial images”“UAV detection and tracking benchmark”“anti-UAV tracking”“visual servoing quadrotor target tracking”“moving target interception UAV”“AirSim autonomous aerial vehicle vision simulation”等。',
        '在文献纳入方面，本文优先保留与低慢小目标检测、无人机/反无人机目标跟踪、视觉伺服与引导控制、仿真验证直接相关的论文；来源上优先选取 SCI 期刊、IEEE/CVF/AAAI/ECCV/ICCV/CVPR/WACV 等高质量会议论文，并适当保留少量高被引基础性方法论文。最终筛得 40 篇文献，其中小目标检测 11 篇，数据集与基准 12 篇，跟踪与状态估计 8 篇，视觉伺服与拦截控制及仿真 9 篇。',
        '需要说明的是，综述的目标并非逐篇罗列论文，而是围绕技术路线展开主题化综合。本文在每一部分均重点分析相关文献的核心思想、适用场景与对本课题的启发，从而为后续论文撰写和实验设计提供直接支持。',
    ]:
        add_body(doc, para)

    add_heading(doc, '3 低慢小目标视觉检测研究现状')
    add_heading(doc, '3.1 通用检测框架对小目标任务的启发', level=2)
    for para in [
        '在通用目标检测领域，FPN[1]通过自顶向下的特征融合机制实现了高层语义与低层细节信息的联合表示，为小目标提供了更强的多尺度特征基础；Focal Loss[2]则通过重加权难样本的方式缓解了前景 - 背景极度不平衡问题，对于低慢小目标这种正样本比例极低的任务具有重要意义。随后，EfficientDet[3]在特征金字塔和网络缩放策略上进一步提升了检测效率与精度平衡，为边缘端部署提供了工程参考。',
        'Transformer 检测器的发展也为该方向带来新的思路。DETR[4]以集合预测方式统一了检测流程，但对小目标和收敛速度仍存在挑战；QueryDet[5]则针对高分辨率图像中的稀疏小目标提出级联稀疏查询机制，强调仅在可能存在目标的位置进行细粒度推理，这种思想对于低空场景中稀疏出现的小型飞行器具有较高借鉴价值。总体来看，通用检测器在特征融合、损失设计和稀疏计算方面提供了小目标检测的基础能力，但其直接应用于低空微小飞行目标时仍需场景化改造。',
    ]:
        add_body(doc, para)
    add_heading(doc, '3.2 面向无人机/航拍场景的小目标检测改进', level=2)
    for para in [
        '针对航拍视角中目标尺度小、背景复杂和目标姿态变化大的问题，研究者开始围绕上下文建模、多尺度融合和注意力机制展开专门改进。文献[6]通过特征融合和尺度增强提升了无人机图像中小目标的定位能力，说明单层特征难以支撑微小目标检测；文献[7]从“先聚焦再检测”的角度降低背景干扰，将检测过程与区域选择耦合；文献[8]进一步采用多尺度特征融合，显示在 UAV 场景下高分辨率浅层纹理信息与深层语义信息必须协同使用。',
        '此外，文献[9]构建了针对无人机图像中“小弱目标”的新基准，并强调上下文尺度敏感建模的重要性；文献[10]利用混合注意力强化目标与背景区分；文献[11]则从 tiny object 的角度总结了在航拍图像中由目标像素占比极低所引出的定位与分类难题。这些工作共同说明，低慢小目标检测并不只是通用检测器的轻量化移植，而需要围绕尺度适配、上下文信息、局部细节恢复和复杂背景抑制进行专门优化。',
    ]:
        add_body(doc, para)
    add_heading(doc, '3.3 对本课题的启示', level=2)
    add_body(doc, '结合本文毕设项目已有实现，YOLO 路线仍然具有较高工程可行性：一方面，YOLO 具备较好的推理实时性，适合与控制环路耦合；另一方面，本项目已经完成仿真采集、自动标注和训练脚本链路，能够较低成本地完成针对低慢小目标的场景化微调。但从文献综述可以看出，若要进一步提升性能，后续可重点考虑多尺度特征融合、注意力增强、小目标样本增广和漏检后跟踪保持四类问题。')

    add_heading(doc, '4 无人机检测跟踪数据集与评测基准')
    for para in [
        '数据集和评测基准是低慢小目标研究走向规范化的重要前提。UAVDT 基准[13]和 Detection and Tracking Meet Drones Challenge[12]较早系统地将无人机视角下的目标检测、单目标跟踪和多目标跟踪任务统一起来，为后续算法比较提供了评价基础。相关挑战和综述类工作表明，无人机场景与地面监控场景存在明显差异，主要体现在视角俯视、目标尺度极小、平台运动带来的背景变化以及长距离拍摄条件下的纹理退化。',
        '围绕反无人机场景，Vision-Based Anti-UAV Detection and Tracking[14]与 Anti-UAV 基准[15]进一步聚焦“无人机目标本身”作为被观测对象的问题，强调小尺度飞行目标在复杂天空或地面背景中的可见性不足和长时遮挡问题。A Unified Transformer-Based Tracker for Anti-UAV Tracking[16]说明，随着反无人机任务的发展，单纯依赖检测或传统相关滤波方法已难以满足复杂场景需求，鲁棒特征建模和时空上下文表示成为重要方向。',
        '在更广泛的无人机视觉基准中，SeaDronesSee[17]面向海上搜救场景，BIRDSAI[18]强调红外热视频中的检测与跟踪，WebUAV-3M[19]则通过大规模单目标跟踪数据进一步推动了深度跟踪模型训练。与此同时，VisDrone 挑战结果文献[21][22]和基准比较工作[20]持续完善了无人机检测与视频理解评测体系。对于本文课题而言，这些数据集与基准至少提供了两方面启示：第一，评价指标不能只看检测精度，还应考虑跟踪连续性和系统级稳定性；第二，若目标是拦截而非离线识别，就需要关注数据集是否包含对控制有意义的时间序列信息。',
    ]:
        add_body(doc, para)

    add_heading(doc, '5 目标跟踪、状态估计与运动关联方法')
    add_heading(doc, '5.1 检测后关联与多目标跟踪', level=2)
    for para in [
        '在检测结果基础上完成目标时序关联，是将视觉检测扩展为动态跟踪系统的关键环节。Deep SORT[24]在 SORT 的基础上引入外观特征关联，使目标在短时遮挡与交叉情况下的 ID 保持能力显著增强，成为检测后跟踪领域的经典基线。ByteTrack[25]进一步提出“将低分检测框也纳入关联”的思想，指出在动态场景中被滤除的低分框往往包含维持轨迹连续性的关键信息，这对于低慢小目标因尺度过小而导致置信度不稳定的情况具有重要参考价值。',
        '进一步地，Observation-Centric SORT[30]重新审视以运动预测为核心的在线跟踪框架，强调观测驱动的关联过程与更鲁棒的状态更新方式。HOTA 指标[31]则从评测层面对检测质量、关联质量和定位质量进行了统一建模，使研究者能够更全面地评价跟踪系统。对本文课题而言，虽然当前任务主要关注单目标拦截，但多目标跟踪中的状态更新和低置信度利用思想同样适用于“检测 - 预测 - 控制”链路的稳定化。',
    ]:
        add_body(doc, para)
    add_heading(doc, '5.2 单目标跟踪、端到端跟踪与状态预测', level=2)
    for para in [
        '在单目标跟踪方面，SiamRPN++[27]和 DiMP[28]分别代表了基于孪生网络与基于判别式模型预测的主流路线。前者强调匹配表示学习和搜索区域内的高效定位，后者则将模型更新与判别性保持结合起来，在复杂背景与外观变化下表现更稳健。CenterTrack/Tracking Objects as Points[26]和 Improving Multiple Object Tracking with Single Object Tracking[29]则从端到端时序建模的角度，将检测和跟踪更紧密地耦合在一起，体现了“从分离式模块转向统一式时序推理”的发展趋势。',
        '不过，对于毕设级闭环拦截系统来说，复杂深度跟踪器未必是最优选择。原因在于，拦截控制更关心目标图像中心、尺度变化和短时运动趋势，而非长时重识别本身。当前项目中引入的 Kalman 预测模块正好位于检测与控制之间：当检测结果可靠时，利用观测更新状态；当出现短时漏检或抖动时，依赖状态预测维持控制输出连续。文献综述表明，这类“轻量运动模型 + 视觉观测校正”的设计，在工程系统中仍具有很高价值。',
    ]:
        add_body(doc, para)

    add_heading(doc, '6 视觉伺服、目标逼近与拦截控制研究现状')
    for para in [
        '视觉伺服研究为从图像观测到飞行控制量的转换提供了理论基础。Position-Based Visual Servoing for Target Tracking by a Quadrotor UAV[35]和 Toward Visibility Guaranteed Visual Servoing Control of Quadrotor UAVs[32]表明，仅有目标检测并不足以支撑稳定的飞行追踪，还必须考虑视场保持、可见性约束、姿态变化和控制稳定性等因素。前者强调目标相对位姿估计与控制律设计，后者则将目标始终保持在可观测范围内作为重要控制目标。',
        '随着视觉感知与控制融合加深，Image-Based Visual Servoing of Rotorcrafts to Planar Visual Targets of Arbitrary Orientation[34]和 Image-Based Visual Servoing of Quadrotors to Arbitrary Flight Targets[33]进一步说明，直接从图像误差构建控制器，能够减少中间状态估计误差对控制的累积影响；Autonomous Target Tracking of UAV Using High-Speed Visual Feedback[37]则强调高刷新视觉反馈在动态目标追踪中的意义。与此同时，Robust Visual Servoing Formation Tracking Control for Quadrotor UAV Team[38]展示了视觉控制在多飞行器协同中的可扩展性。',
        '更接近拦截场景的工作包括 Autonomous Drone Hunter Operating by Deep Learning and All-Onboard Computations in GPS-Denied Environments[36] 和 UAV Guidance: A Stereo-Based Technique for Interception of Stationary or Moving Targets[39]。这些研究证明，在 GPS 受限或环境复杂场景下，依靠视觉感知、目标定位和自主控制可以完成对空中目标的持续追踪甚至捕获。与这些工作相比，本文毕设选择的路线更强调工程可复现性：利用图像中心误差、面积比、短时预测和状态机控制完成搜索、跟踪、逼近与捕获，属于一种更轻量、更适合本科毕设实现条件的视觉引导拦截方案。',
    ]:
        add_body(doc, para)

    add_heading(doc, '7 仿真平台、系统集成与本文课题切入点')
    for para in [
        '仿真平台是将检测、跟踪、预测和控制整合为闭环系统的重要支撑。AirSim[40]提供了较高保真的视觉与物理仿真能力，使研究者能够在无需昂贵实机试验成本的前提下，完成场景复现、数据采集、策略调试与性能对比。在本文对应的毕设项目中，Collect.py 已实现图像采集与自动标注流程，Train.py 完成检测模型训练，run.py 将 YOLO 检测、BBox 预测与引导控制串联起来，VisualInterceptController 则在 UE 侧承担状态机与控制分发功能，构成了完整的闭环原型。',
        '从现有文献来看，当前研究仍存在若干不足。第一，复杂背景下低慢小目标仍容易漏检，检测器的时序稳定性不足；第二，多数基准以检测或跟踪任务为主，较少直接面向“感知 - 预测 - 控制 - 拦截成功率”的系统级评价；第三，很多工作将检测、跟踪和控制割裂处理，导致漏检后的控制退化问题没有得到充分解决；第四，部分高性能方法依赖较重模型或多传感器条件，难以直接迁移到低成本、轻量化闭环实验平台。',
        '因此，本文课题的研究切入点可以明确概括为：依托仿真环境构建低慢小目标数据集；在 YOLO 检测模型基础上完成面向小目标的场景化适配；通过卡尔曼预测增强视觉观测连续性；以图像中心误差、目标面积和状态机为核心构建视觉引导拦截控制律；最终以检测效果、控制稳定性和拦截成功率为指标完成系统验证。这一方案既与文献发展趋势一致，也与当前项目代码基础高度契合。',
    ]:
        add_body(doc, para)

    add_heading(doc, '8 结论')
    for para in [
        '综上所述，低慢小目标视觉拦截研究已经形成较清晰的技术演进脉络：在感知层面，研究由通用检测框架逐步发展到面向航拍场景的小目标专用模型；在数据层面，研究由零散实验走向 UAVDT、VisDrone、Anti-UAV、WebUAV-3M 等标准化基准；在时序层面，研究由简单卡尔曼关联逐步演化到融合检测、外观、运动和端到端时序建模的复杂跟踪方法；在控制层面，研究由传统视觉伺服逐步迈向面向动态目标的自主追踪与拦截控制。',
        '对于本文毕设而言，最有价值的结论并不是简单追求单项精度最优，而是在有限工程条件下构建一条稳定、闭环、可复现的系统链路。现有文献为检测器改进、基准选择、跟踪与预测建模、视觉控制设计和仿真验证提供了充分依据。后续写作中，可围绕“低慢小目标视觉检测优化、漏检下的目标状态预测、视觉引导拦截控制策略、系统级实验评估”四个核心问题展开论文正文，从而形成与项目实现相一致、结构完整且逻辑清晰的毕业论文。',
    ]:
        add_body(doc, para)

    add_heading(doc, '参考文献')
    for ref in references:
        add_reference(doc, ref['id'], f"{ref['title']}. {ref['venue']}, {ref['year']}. DOI: {ref['doi']}")

    doc.save(REVIEW_PATH)


if __name__ == '__main__':
    build_task_doc()
    build_review_doc()
    print(TASK_PATH)
    print(REVIEW_PATH)
